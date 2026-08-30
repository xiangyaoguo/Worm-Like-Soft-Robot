from __future__ import annotations

import argparse
import csv
import json
import math
import os
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


SOURCE_STUDY_ROOT = Path(__file__).resolve().parents[1]
STUDY_ROOT = Path(
    os.environ.get("THESIS_HPR_OUTPUT", str(SOURCE_STUDY_ROOT))
).resolve()
DATA_ROOT = STUDY_ROOT / "data"
FIGURE_ROOT = STUDY_ROOT / "figures"
REPORT_FIGURE_ROOT = FIGURE_ROOT / "report_jpg"
REPORT_ROOT = STUDY_ROOT / "report"
OUTPUT_PATH = REPORT_ROOT / "Formal_HPR_Rolling_Mechanism_Validation_English_20260810.docx"

from table_geometry import apply_table_geometry, column_widths_from_weights


BLUE = "1F5A91"
DARK_BLUE = "172B4D"
LIGHT_BLUE = "EAF2F8"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "667085"
WHITE = "FFFFFF"
RED = "B42318"
GREEN = "157A6E"
GOLD = "B54708"
TOTAL_TABLE_WIDTH = 9360


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def load_json(path: Path) -> Any:
    return json.loads(path.read_text(encoding="utf-8-sig"))


def set_run_font(
    run,
    name: str = "Calibri",
    east_asia: str = "Calibri",
    size: float | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    colour: str | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if colour is not None:
        run.font.color.rgb = RGBColor.from_string(colour)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instruction, separate, text, end):
        run._r.append(element)
    set_run_font(run, size=9, colour=MID_GRAY)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, before, after, colour in (
        ("Heading 1", 16, 16, 8, BLUE),
        ("Heading 2", 13, 12, 6, BLUE),
        ("Heading 3", 11.5, 8, 4, DARK_BLUE),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(colour)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = doc.styles["Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    caption.font.size = Pt(9)
    caption.font.italic = False
    caption.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.keep_with_next = False

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Cm(1.27)
        style.paragraph_format.first_line_indent = Cm(-0.635)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.10


def configure_section(section, first_page: bool = False) -> None:
    # Named override to standard_business_brief: A4 page geometry for a UK MSc technical report.
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    if first_page:
        section.different_first_page_header_footer = True


def configure_running_furniture(section) -> None:
    header = section.header
    header.is_linked_to_previous = False
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("FORMAL HPR MECHANISM VALIDATION | EVIDENCE REPORT")
    set_run_font(run, size=8.5, bold=True, colour=MID_GRAY)

    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_before = Pt(0)
    run = paragraph.add_run("10 August 2026  |  ")
    set_run_font(run, size=8.5, colour=MID_GRAY)
    add_page_number(paragraph)


def add_report_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading(level=level)
    paragraph.add_run(text)


def add_report_paragraph(
    doc: Document,
    text: str,
    *,
    bold_lead: str | None = None,
) -> None:
    paragraph = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_run_font(lead, bold=True, colour=DARK_BLUE)
        paragraph.add_run(text[len(bold_lead) :])
    else:
        paragraph.add_run(text)


def add_callout(doc: Document, text: str, colour: str = BLUE) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.35)
    paragraph.paragraph_format.right_indent = Cm(0.35)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.keep_together = True
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), LIGHT_BLUE)
    p_pr.append(shading)
    border = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "5")
    left.set(qn("w:color"), colour)
    border.append(left)
    p_pr.append(border)
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5, bold=True, colour=DARK_BLUE)


def add_table(
    doc: Document,
    headers: list[str],
    rows: Iterable[Iterable[str]],
    widths: list[float],
    caption_text: str,
) -> None:
    caption = doc.add_paragraph(style="Caption")
    caption.paragraph_format.keep_with_next = True
    run = caption.add_run(caption_text)
    set_run_font(run, size=9, bold=True, colour=DARK_BLUE)
    row_values = [list(row) for row in rows]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.rows[0].cells
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cell, LIGHT_GRAY)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            set_run_font(run, size=8.5, bold=True, colour=DARK_BLUE)
    set_repeat_table_header(table.rows[0])
    for row_index, values in enumerate(row_values):
        cells = table.add_row().cells
        for column_index, (cell, value) in enumerate(zip(cells, values)):
            cell.text = str(value)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT if column_index == 0 else WD_ALIGN_PARAGRAPH.CENTER
            )
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            for run in paragraph.runs:
                set_run_font(run, size=8.2, colour=DARK_BLUE)
        if row_index % 2 == 1:
            for cell in cells:
                set_cell_shading(cell, "FAFBFC")
    column_widths = column_widths_from_weights(widths, TOTAL_TABLE_WIDTH)
    apply_table_geometry(
        table,
        column_widths,
        table_width_dxa=TOTAL_TABLE_WIDTH,
        indent_dxa=120,
        cell_margins_dxa={"top": 90, "bottom": 90, "start": 120, "end": 120},
    )
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_figure(
    doc: Document,
    figure_path: Path,
    caption_text: str,
    width_inches: float = 6.45,
) -> None:
    if not figure_path.is_file():
        raise FileNotFoundError(figure_path)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    inline_shape = run.add_picture(str(figure_path), width=Inches(width_inches))
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("descr", caption_text)
    caption = doc.add_paragraph(style="Caption")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run(caption_text)
    set_run_font(run, size=9, bold=True, colour=DARK_BLUE)


def row_lookup(rows: list[dict[str, str]], seed: int, condition: str) -> dict[str, str]:
    for row in rows:
        condition_id = row.get("condition_id", row.get("id"))
        if int(row["training_seed"]) == seed and condition_id == condition:
            return row
    raise KeyError((seed, condition))


def success(row: dict[str, str]) -> int:
    return int(float(row["kinematic_success_count"]))


def f(row: dict[str, str], key: str) -> float:
    return float(row[key])


def effect_lookup(effects: list[dict[str, str]], seed: int, condition: str) -> dict[str, str]:
    for row in effects:
        if int(row["training_seed"]) == seed and row["condition_id"] == condition:
            return row
    raise KeyError((seed, condition))


def strongest_joint(effects: list[dict[str, str]], seed: int, suffix: str) -> tuple[int, float]:
    candidates = []
    for joint in range(1, 9):
        row = effect_lookup(effects, seed, f"J{joint:02d}_{suffix}")
        candidates.append((joint, float(row["desired_revolutions_delta_mean"])))
    return min(candidates, key=lambda item: item[1])


def joint_counts(rows: list[dict[str, str]], seed: int, suffix: str) -> list[int]:
    return [success(row_lookup(rows, seed, f"J{joint:02d}_{suffix}")) for joint in range(1, 9)]


def configure_title_page(doc: Document) -> None:
    section = doc.sections[0]
    configure_section(section, first_page=True)
    for _ in range(4):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(10)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.add_run("FORMAL CHECKPOINT INTERVENTION STUDY")
    set_run_font(run, size=10.5, bold=True, colour=BLUE)
    kicker.paragraph_format.space_after = Pt(16)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("Formal HPR Rolling-Mechanism Validation")
    set_run_font(run, size=27, bold=True, colour=DARK_BLUE)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(14)
    run = subtitle.add_run("Formal HPR Rolling-Mechanism Validation")
    set_run_font(run, size=20, bold=True, colour=BLUE)

    descriptor = doc.add_paragraph()
    descriptor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    descriptor.paragraph_format.space_after = Pt(36)
    run = descriptor.add_run(
        "Paired K₁/K₂ and joint-freezing interventions on two independent HPR training runs initialised with random seeds 9201 and 9205"
    )
    set_run_font(run, size=12, colour=MID_GRAY)

    metadata = doc.add_paragraph()
    metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
    metadata.paragraph_format.space_after = Pt(4)
    run = metadata.add_run("Study ID: formal_hpr_freeze_validation_20260810")
    set_run_font(run, size=10, bold=True, colour=DARK_BLUE)
    metadata = doc.add_paragraph()
    metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = metadata.add_run("Evaluation-only analysis | No retraining | 10 August 2026")
    set_run_font(run, size=9.5, colour=MID_GRAY)

    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(section)
    configure_running_furniture(section)
    page_num_type = OxmlElement("w:pgNumType")
    page_num_type.set(qn("w:start"), "1")
    section._sectPr.append(page_num_type)


def build_report() -> Path:
    REPORT_ROOT.mkdir(parents=True, exist_ok=True)
    summary_rows = read_csv(DATA_ROOT / "condition_policy_summary.csv")
    effects = read_csv(DATA_ROOT / "paired_effects.csv")
    analysis = load_json(DATA_ROOT / "ANALYSIS_SUMMARY.json")
    study_manifest = load_json(DATA_ROOT / "STUDY_MANIFEST.json")

    global_rows = {
        (seed, condition): row_lookup(summary_rows, seed, condition)
        for seed in (9201, 9205)
        for condition in ("BASELINE", "GLOBAL_K1_OFF", "GLOBAL_K2_OFF", "GLOBAL_BOTH_OFF")
    }
    strongest = {
        (seed, suffix): strongest_joint(effects, seed, suffix)
        for seed in (9201, 9205)
        for suffix in ("BOTH_OFF", "K1_OFF", "K2_OFF")
    }
    only_counts = {seed: joint_counts(summary_rows, seed, "ONLY") for seed in (9201, 9205)}
    both_off_counts = {seed: joint_counts(summary_rows, seed, "BOTH_OFF") for seed in (9201, 9205)}
    k1_off_counts = {seed: joint_counts(summary_rows, seed, "K1_OFF") for seed in (9201, 9205)}
    k2_off_counts = {seed: joint_counts(summary_rows, seed, "K2_OFF") for seed in (9201, 9205)}
    shared_whole_zero = [
        joint
        for joint in range(1, 9)
        if both_off_counts[9201][joint - 1] == 0
        and both_off_counts[9205][joint - 1] == 0
    ]
    shared_k1_severe = [
        joint
        for joint in range(1, 9)
        if k1_off_counts[9201][joint - 1] <= 4
        and k1_off_counts[9205][joint - 1] <= 4
    ]
    shared_whole_zero_text = ", ".join(f"J{joint:02d}" for joint in shared_whole_zero) or "none"
    shared_k1_severe_text = ", ".join(f"J{joint:02d}" for joint in shared_k1_severe) or "none"

    doc = Document()
    configure_styles(doc)
    configure_title_page(doc)

    add_report_heading(doc, "Executive summary",  1)
    add_report_paragraph(
        doc,
        "This evaluation-only study tested whether the feedback-channel and multi-joint dependencies observed in an earlier selected-policy case study were reproduced in both consistently rolling formal horizontal-progress reward (HPR) policies. The policies at training checkpoint 1500 from two independent training runs initialised with random seeds 9201 and 9205 were evaluated under 36 pre-specified frozen-policy interventions, 20 paired formal reset states and 1,000 control steps, giving 1,440 rollouts.",
        
    )
    k1_9201 = success(global_rows[(9201, "GLOBAL_K1_OFF")])
    k1_9205 = success(global_rows[(9205, "GLOBAL_K1_OFF")])
    k2_9201 = success(global_rows[(9201, "GLOBAL_K2_OFF")])
    k2_9205 = success(global_rows[(9205, "GLOBAL_K2_OFF")])
    both_9201 = success(global_rows[(9201, "GLOBAL_BOTH_OFF")])
    both_9205 = success(global_rows[(9205, "GLOBAL_BOTH_OFF")])
    add_callout(
        doc,
        f"Primary result: baseline rolling was reproduced exactly at 20/20 for both formal policies. Disabling all K₁ outputs changed success to {k1_9201}/20 and {k1_9205}/20; disabling all K₂ outputs changed it to {k2_9201}/20 and {k2_9205}/20; disabling both channels gave {both_9201}/20 and {both_9205}/20.",
        
    )
    max_only_9201 = max(only_counts[9201])
    max_only_9205 = max(only_counts[9205])
    add_report_paragraph(
        doc,
        f"Joint-level effects were heterogeneous across the two learned policies, but retaining any one joint never reproduced baseline robustness: the best single-joint-retention condition reached {max_only_9201}/20 for HPR policy 9201 and {max_only_9205}/20 for HPR policy 9205. The defensible conclusion is therefore that the intervention pattern was replicated across both stable formal HPR rolling policies and is consistent with distributed, closed-loop coordination. It is not evidence for a universal HPR mechanism because only two independent consistently rolling training runs were available.",
        
    )

    doc.add_page_break()
    add_report_heading(doc, "1. Scope and evidence boundary",  1)
    add_report_paragraph(
        doc,
        "The formal HPR experiment used a conventional horizontal-progress objective (internal run label R0). The present work did not alter the reward, observations, actor-network parameters, shared centralised critic, physics or terrain. It intervened only on the deterministic outputs of the frozen joint-specific actor networks immediately before each environment step. Consequently, the study asks which learned feedback channels and joints were functionally required or sufficient for these two fixed rolling policies under matched states.",
        
    )
    add_report_paragraph(
        doc,
        "The 20 reset states are repeated paired conditions nested within each frozen policy. They quantify within-policy robustness and paired intervention effects, but they are not 20 independent training runs. Cross-training replication is supported only by agreement between the policies obtained from the two independent training runs.",
        
    )

    add_report_heading(doc, "2. Methods",  1)
    add_report_heading(doc, "2.1 Formal training checkpoints and controller",  2)
    add_report_paragraph(
        doc,
        "Each ten-particle robot had eight actuated joints. PPO used independently parameterised joint-specific actors and a shared centralised critic. At each joint, the local observation contained the spatial angular difference and the current joint angular velocity; the actor network output the spatial-difference feedback gain, K₁, and angular-velocity feedback gain, K₂. The environment applied the learned gains through the original clipped torque law, while passive elasticity, damping and ground contact remained active in every intervention condition.",
        
    )
    add_table(
        doc,
        ["Training run", "Reward", "Endpoint", "Checkpoint SHA-256"],
        [
            ["run (random seed 9201)", "HPR", "training checkpoint 1500", "90f37918…d3a0ae52"],
            ["run (random seed 9205)", "HPR", "training checkpoint 1500", "3c10cbb2…994aec7e"],
        ],
        [1.1, 0.8, 1.4, 3.2],
        "Table 1. Frozen formal-policy objects.",
        
    )

    add_report_heading(doc, "2.2 Paired intervention design",  2)
    add_report_paragraph(
        doc,
        "The full matrix comprised the unmodified baseline; global K₁-off, K₂-off and both-off conditions; eight joint-specific K₁ ablations; eight joint-specific K₂ ablations; eight whole-joint ablations; and eight single-joint-retention conditions. For each policy-condition pair, reset seeds 20264101–20264120 were replayed deterministically for 1,000 control steps. Disabling a learned channel set only the corresponding actor-network output to zero at every step; it did not remove passive mechanics or freeze the robot configuration.",
        
    )
    add_figure(
        doc,
        REPORT_FIGURE_ROOT / "fig_01_intervention_design.jpg",
        "Figure 1. Formal HPR frozen-policy intervention design and evidence hierarchy.",
        
        width_inches=5.75,
    )

    add_report_heading(doc, "2.3 Rolling endpoint",  2)
    add_report_paragraph(
        doc,
        "A rollout satisfied the common kinematic full-body rolling criterion only when desired-direction net best-fit body rotation was at least 360°, the desired active-rotation fraction was at least 0.70, and forward displacement was at least one initial body length. This criterion was selected because the HPR environment did not export the later contact/support event fields. Pulse counts were retained only as diagnostics and were not imputed into the primary endpoint.",
        
    )

    add_report_heading(doc, "3. Integrity and baseline verification",  1)
    baseline_rows = []
    for seed in (9201, 9205):
        row = global_rows[(seed, "BASELINE")]
        identity = load_json(DATA_ROOT / "raw" / f"seed{seed}" / "BASELINE_IDENTITY_AUDIT.json")
        max_error = max(float(value) for value in identity["maximum_absolute_error"].values())
        baseline_rows.append(
            [
                str(seed),
                f"{success(row)}/20",
                f"{f(row, 'desired_revolutions_mean'):.3f} ± {f(row, 'desired_revolutions_sample_sd'):.3f}",
                f"{f(row, 'direction_fraction_mean'):.3f}",
                f"{f(row, 'forward_body_lengths_mean'):.3f}",
                f"{max_error:.1e}",
            ]
        )
    add_table(
        doc,
        ["HPR policy", "Success", "Desired rev. mean ± SD", "Direction fraction", "Forward BL", "Max identity error"],
        baseline_rows,
        [0.7, 0.7, 1.8, 1.2, 1.0, 1.1],
        "Table 2. Exact baseline reproduction against the formal endpoint evaluation.",
        
    )
    add_report_paragraph(
        doc,
        "Both baselines reproduced all 20 formal reset trajectories. For the six audited per-episode metrics—initial body length, displacement, displacement in body lengths, net rotation, desired net rotation and directional fraction—the maximum absolute discrepancy from the stored formal evaluation was zero. Training-checkpoint and in-memory policy hashes were unchanged after the full matrix.",
        
    )

    add_report_heading(doc, "4. Results",  1)
    add_report_heading(doc, "4.1 Global feedback-channel interventions",  2)
    global_table_rows = []
    for seed in (9201, 9205):
        for condition in ("BASELINE", "GLOBAL_K1_OFF", "GLOBAL_K2_OFF", "GLOBAL_BOTH_OFF"):
            row = global_rows[(seed, condition)]
            global_table_rows.append(
                [
                    str(seed),
                    {
                        "BASELINE": "Baseline",
                        "GLOBAL_K1_OFF": "K₁ off",
                        "GLOBAL_K2_OFF": "K₂ off",
                        "GLOBAL_BOTH_OFF": "Both off",
                    }[condition],
                    f"{success(row)}/20",
                    f"{f(row, 'desired_revolutions_mean'):.3f}",
                    f"{f(row, 'direction_fraction_mean'):.3f}",
                    f"{f(row, 'forward_body_lengths_mean'):.3f}",
                ]
            )
    add_table(
        doc,
        ["HPR policy", "Condition", "Success", "Desired rev.", "Direction", "Forward BL"],
        global_table_rows,
        [0.7, 1.55, 0.8, 1.1, 1.0, 1.0],
        "Table 3. Global learned-feedback channel outcomes.",
        
    )
    add_report_paragraph(
        doc,
        f"The global intervention pattern was evaluated independently for each formal policy. K₁-off yielded {k1_9201}/20 and {k1_9205}/20 successes; K₂-off yielded {k2_9201}/20 and {k2_9205}/20; both-off yielded {both_9201}/20 and {both_9205}/20. Continuous rotation, directional consistency and displacement are reported alongside binary success because some conditions approached individual thresholds and because binary counts alone conceal the type of failure.",
        
    )
    add_figure(
        doc,
        REPORT_FIGURE_ROOT / "fig_02_global_channel_outcomes.jpg",
        "Figure 2. Global channel interventions. Error bars show the sample standard deviation across 20 nested paired reset states; they are not uncertainty across independent training runs.",
        
    )

    add_report_heading(doc, "4.2 Whole-joint ablation and single-joint retention",  2)
    j9201, d9201 = strongest[(9201, "BOTH_OFF")]
    j9205, d9205 = strongest[(9205, "BOTH_OFF")]
    add_report_paragraph(
        doc,
        f"Removing both learned outputs from one joint revealed both shared and policy-dependent structure. Disabling {shared_whole_zero_text} individually eliminated all 20 successes in both policies. The largest paired reduction in desired revolutions occurred at J{j9201:02d} for HPR policy 9201 ({d9201:+.2f} revolutions) and J{j9205:02d} for HPR policy 9205 ({d9205:+.2f} revolutions). The complete J01–J08 success-count profiles were {both_off_counts[9201]} and {both_off_counts[9205]}, respectively. Thus, several anterior joints showed reproducible necessity, while the remaining ranking was policy dependent.",
        
    )
    add_report_paragraph(
        doc,
        f"Single-joint retention was a stricter sufficiency test. Across J01–J08, the success-count profiles were {only_counts[9201]} for HPR policy 9201 and {only_counts[9205]} for HPR policy 9205. No single retained joint matched the 20/20 baseline in either policy; therefore the data reject a single-joint-sufficient explanation for robust full-body rolling in these formal policies.",
        
    )
    add_figure(
        doc,
        REPORT_FIGURE_ROOT / "fig_03_joint_ablation_retention_heatmap.jpg",
        "Figure 3. Whole-joint ablation and single-joint retention. The upper panel reports success rates; the lower panel reports paired change in desired revolutions relative to the matched baseline.",
        
    )

    add_report_heading(doc, "4.3 Joint-specific K₁ and K₂ contributions",  2)
    k1j9201, k1d9201 = strongest[(9201, "K1_OFF")]
    k1j9205, k1d9205 = strongest[(9205, "K1_OFF")]
    k2j9201, k2d9201 = strongest[(9201, "K2_OFF")]
    k2j9205, k2d9205 = strongest[(9205, "K2_OFF")]
    add_report_paragraph(
        doc,
        f"Channel-separated ablations further showed that joint importance depended on both location and learned policy. K₁ ablation at {shared_k1_severe_text} reduced success to at most 4/20 in both policies, providing the strongest shared channel-specific evidence. The largest K₁-related revolution reductions occurred at J{k1j9201:02d} ({k1d9201:+.2f}) and J{k1j9205:02d} ({k1d9205:+.2f}) for HPR policies 9201 and 9205. The strongest K₂-related reductions occurred at J{k2j9201:02d} ({k2d9201:+.2f}) and J{k2j9205:02d} ({k2d9205:+.2f}). The remaining differences should be interpreted as policy-specific functional effects rather than fixed biological or mechanical constants.",
        
    )
    add_figure(
        doc,
        REPORT_FIGURE_ROOT / "fig_04_joint_channel_effect_heatmap.jpg",
        "Figure 4. Joint-specific K₁ and K₂ ablations across both formal HPR policies.",
        
    )

    add_report_heading(doc, "4.4 Cross-policy agreement",  2)
    agreement = analysis["cross_seed_agreement"]
    whole_rho = agreement["whole_joint_ablation"]["spearman_rank_correlation"]
    k1_rho = agreement["joint_k1_ablation"]["spearman_rank_correlation"]
    k2_rho = agreement["joint_k2_ablation"]["spearman_rank_correlation"]
    add_report_paragraph(
        doc,
        f"The descriptive cross-policy rank correlations of paired revolution effects were {whole_rho:.2f} for whole-joint ablation, {k1_rho:.2f} for K₁ ablation and {k2_rho:.2f} for K₂ ablation. These values quantify whether the two learned policies ranked joint effects similarly; they do not constitute a population-level significance test with only two independent policies.",
        
    )
    add_figure(
        doc,
        REPORT_FIGURE_ROOT / "fig_05_cross_seed_agreement.jpg",
        "Figure 5. Descriptive cross-policy agreement in paired joint-intervention effects.",
        
        width_inches=5.9,
    )

    add_report_heading(doc, "5. Representative morphology and temporal evidence",  1)
    add_report_paragraph(
        doc,
        "Figures 6–7 use the same formal reset seed, 20264101, for every displayed condition. The morphology panels centre each body for shape comparison and annotate centre-of-mass displacement in body lengths. They illustrate how the global interventions altered the closed-loop trajectory, but the statistical summaries above use all 20 paired reset states. A matching cumulative rotation and displacement plot was also generated as a supplementary standalone asset.",
        
    )
    add_figure(
        doc,
        REPORT_FIGURE_ROOT / "fig_06_morphology_seed9201.jpg",
        "Figure 6. Matched morphology sequence for formal HPR policy 9201.",
        
    )
    add_figure(
        doc,
        REPORT_FIGURE_ROOT / "fig_07_morphology_seed9205.jpg",
        "Figure 7. Matched morphology sequence for formal HPR policy 9205.",
        
    )
    add_report_heading(doc, "6. Interpretation and thesis claim",  1)
    add_report_paragraph(
        doc,
        "The strongest supported claim is that robust full-body rolling in both consistently rolling formal HPR policies depended on the learned K₁/K₂ outputs and multi-joint participation, and was consistent with a distributed closed loop mediated by body mechanics and ground contact. Global interventions tested channel-level dependence, whereas joint ablation and retention demonstrated that the system could not be reduced to one independently sufficient joint. The policy-dependent joint rankings are consistent with more than one local coordination pattern under the same horizontal-progress objective.",
        
    )
    add_callout(
        doc,
        "Claim boundary: write ‘reproduced in both consistently rolling formal HPR policies’, not ‘proved for all HPR policies’.",
        
        colour=GOLD,
    )

    doc.add_page_break()
    add_report_heading(doc, "7. Conclusion",  1)
    add_report_paragraph(
        doc,
        "The requested formal replication was completed without retraining or training-checkpoint modification. Both formal HPR rolling policies showed a common strong dependence on K₁ and on multi-joint participation, while the magnitude of K₂ and joint-specific effects differed between policies. The new evidence therefore upgrades the earlier c10 K₁ and multi-joint dependency observation from a single selected-policy case study to a two-policy formal training-checkpoint replication, while retaining an explicit two-training-run limitation.",
        
    )

    add_report_heading(doc, "8. Copy-ready thesis paragraph",  1)
    thesis_text = (
        f"To test whether the selected-policy mechanism generalised beyond the original case study, frozen-policy interventions were repeated on two consistently rolling formal HPR policies obtained from independent training runs initialised with random seeds 9201 and 9205 at training checkpoint 1500. Each condition used the same 20 deterministic 1,000-step reset states. The unmodified policies exactly reproduced the formal baseline at 20/20 kinematic rolling successes. Disabling the spatial-difference feedback gain, K₁, at all joints reduced success to {k1_9201}/20 and {k1_9205}/20, whereas disabling the angular-velocity feedback gain, K₂, reduced it to {k2_9201}/20 and {k2_9205}/20; disabling both learned channels yielded {both_9201}/20 and {both_9205}/20. Whole-joint ablation of {shared_whole_zero_text} individually eliminated all successes in both policies, and K₁ ablation at {shared_k1_severe_text} reduced both policies to at most 4/20. Retaining any single joint failed to recover the 20/20 baseline (maximum {max_only_9201}/20 and {max_only_9205}/20, respectively), while the remaining joint effects produced heterogeneous rankings across policies. These results reproduce the channel-level and multi-joint dependency pattern in both stable formal HPR rolling policies and support a distributed body–contact feedback interpretation. Because only two independent rolling training runs were available, the evidence does not establish a universal joint hierarchy or a mechanism applicable to every HPR-trained policy."
    )
    add_report_paragraph(doc, thesis_text)

    doc.add_page_break()
    add_report_heading(doc, "9. Limitations and remaining experiments",  1)
    limitations = [
        (
            "Training-run scope",
            "Only two independent formal HPR policies consistently rolled; additional HPR training runs are required for population-level inference.",
        ),
        (
            "Endpoint scope",
            "The primary endpoint is kinematic because HPR did not export the later support-event fields; contact-aware pulse claims require a common re-evaluation environment.",
        ),
        (
            "Physical scope",
            "Results are limited to two-dimensional flat-ground simulation with the locked ten-particle model and do not establish hardware or multi-terrain generalisation.",
        ),
        (
            "Mechanistic scope",
            "Post-policy zeroing is a causal intervention on these frozen closed loops, but it does not identify a unique minimal sufficient joint subset or isolate substep-level energy transfer.",
        ),
    ]
    for label, text in limitations:
        add_report_paragraph(
            doc,
            f"{label}: {text}",
            
            bold_lead=f"{label}: ",
            
        )

    add_report_heading(doc, "10. Reproducibility index",  1)
    add_report_paragraph(
        doc,
        "The source training checkpoints and frozen formal runtime were read only. The study directory contains the preregistered protocol, dedicated runner, condition-level JSON records, per-rollout and summary CSV files, paired effects, trajectory archives, all figures and this report. The study manifest records 1,440 completed rollouts and the post-run integrity receipts for both policies.",
        
    )
    add_table(
        doc,
        ["Artifact", "Location / status"],
        [
            ["Protocol", "PREREGISTERED_PROTOCOL.md"],
            ["Runner", r"scripts\run_formal_hpr_freeze_study.py"],
            ["Per-rollout data", r"data\episode_results.csv"],
            ["Condition summaries", r"data\condition_policy_summary.csv"],
            ["Paired effects", r"data\paired_effects.csv"],
            ["Study manifest", f"{study_manifest['total_rollouts']} rollouts; integrity passed"],
            ["Figures", "figures/"],
        ],
        [1.45, 5.05],
        "Table 4. Reproducibility artifacts.",
        
    )

    doc.core_properties.title = "Formal HPR Rolling-Mechanism Validation"
    doc.core_properties.subject = "Paired K1/K2 and joint-freezing interventions"
    doc.core_properties.author = "Formal HPR validation study"
    doc.core_properties.keywords = "HPR; reinforcement learning; rolling; K1; K2; frozen-policy intervention"
    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Build the English-only Formal HPR rolling-mechanism validation report."
    )
    parser.parse_args()
    print(build_report())


if __name__ == "__main__":
    main()
