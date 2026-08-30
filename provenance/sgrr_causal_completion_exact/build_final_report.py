from __future__ import annotations

"""Build the final English K1/K2 causal-mechanism report.

This builder is deliberately fail-closed.  It validates the three completed
analysis gates before reading any derived result, then assembles a compact
25--40 page narrative report from the frozen outputs.  It never edits training
artifacts, checkpoints, traces, or analysis tables.
"""

import argparse
import hashlib
import json
import math
import os
import re
import tempfile
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterable, Mapping, Sequence

import pandas as pd
from PIL import Image

# `--validate-only` is intentionally usable in the frozen analysis runtime,
# which may not carry python-docx.  DOCX symbols are required only after all
# evidence gates pass and a real report build is requested.
_DOCX_IMPORT_ERROR: ModuleNotFoundError | None = None
try:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
except ModuleNotFoundError as error:
    _DOCX_IMPORT_ERROR = error


ROOT = Path(__file__).resolve().parent
PROJECT_ROOT = ROOT.parent
CAUSAL_DIR = ROOT / "analysis_causal_completion"
LOCAL_DIR = ROOT / "analysis_local_actor"
PROPAGATION_DIR = ROOT / "analysis_physical_propagation"
MATCHED_C00_ROOT = ROOT / "matched_c00"
OLD_MECHANISM_ROOT = PROJECT_ROOT / "mechanism_runtime_exact"
ATLAS_ROOT = (
    PROJECT_ROOT
    / "05_Experimental_Data_and_Code"
    / "formal10_initial3_k1k2_atlas_20260804"
)
DEFAULT_OUTPUT = ROOT / "K1K2_Per_Joint_Causal_Mechanism_and_Rolling_Interpretation_Report_20260804.docx"
DEFAULT_MANIFEST = ROOT / "FINAL_REPORT_MANIFEST.json"

JOINTS = [f"J{i:02d}" for i in range(1, 9)]
CHANNELS = ["K1", "K2"]
EVENT_ORDER = [
    "official_prelaunch",
    "official_pulse_q1",
    "official_pulse_q2",
    "official_pulse_q3",
    "official_pulse_q4",
    "official_pulse_q5",
    "official_rolling_outside_pulse",
]
EVENT_CN = {
    "official_prelaunch": "Prelaunch (no samples in these data)",
    "official_pulse_q1": "Pulse Q1",
    "official_pulse_q2": "Pulse Q2",
    "official_pulse_q3": "Pulse Q3",
    "official_pulse_q4": "Pulse Q4",
    "official_pulse_q5": "Pulse Q5",
    "official_rolling_outside_pulse": "Rolling outside pulses",
}
CLASS_CN = {
    "strong_necessity": "strongly necessary",
    "necessary_contribution": "necessary contribution",
    "timing_critical": "timing-critical",
    "single_channel_sufficient_in_legacy_C00_background": "single-channel sufficient on legacy C00 background",
    "equivalent_or_redundant_when_zeroed_in_C11": "approximately equivalent/redundant when zeroed in C11",
    "no_frozen_threshold_label": "no frozen criterion met",
}

# narrative_proposal preset, resolved to exact tokens.
PRESET = {
    "name": "narrative_proposal",
    "page_width_in": 8.5,
    "page_height_in": 11.0,
    "margin_in": 1.0,
    "header_footer_distance_in": 0.492,
    "usable_width_dxa": 9360,
    "table_indent_dxa": 120,
    "body_font": "Calibri",
    "east_asia_font": "Microsoft YaHei",
    "body_size_pt": 11.0,
    "body_after_pt": 8.0,
    "body_line_twips": 320,
    "h1_size_pt": 16.0,
    "h1_color": "2E74B5",
    "h1_before_pt": 18.0,
    "h1_after_pt": 10.0,
    "h2_size_pt": 13.0,
    "h2_color": "2E74B5",
    "h2_before_pt": 12.0,
    "h2_after_pt": 6.0,
    "h3_size_pt": 12.0,
    "h3_color": "1F4D78",
    "h3_before_pt": 8.0,
    "h3_after_pt": 4.0,
    "list_marker_dxa": 261,
    "list_text_dxa": 540,
    "list_hanging_dxa": 279,
    "list_after_twips": 80,
    "list_line_twips": 290,
    "table_header_fill": "F4F6F9",
    "cell_margins_dxa": {"top": 80, "bottom": 80, "start": 120, "end": 120},
    "navy": "203748",
    "blue": "2E74B5",
    "dark_blue": "1F4D78",
    "gold": "A77A19",
    "muted": "596775",
    "ink": "17212B",
    "light_blue": "EEF4F8",
    "light_gold": "FFF6DC",
    "light_red": "FDEEEE",
    "green": "20734B",
    "red": "9B1C1C",
}


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def load_json(path: Path) -> dict[str, Any]:
    with path.open("r", encoding="utf-8-sig") as handle:
        value = json.load(handle)
    if not isinstance(value, dict):
        raise RuntimeError(f"Expected JSON object: {path}")
    return value


def read_csv(path: Path) -> pd.DataFrame:
    return pd.read_csv(path, encoding="utf-8-sig")


def first_existing(*paths: Path) -> Path | None:
    return next((path for path in paths if path.is_file()), None)


def require_files(paths: Iterable[Path]) -> None:
    missing = [str(path) for path in paths if not path.is_file()]
    if missing:
        raise RuntimeError(
            "The final report refuses incomplete results; the following completion evidence or data files are missing:\n- "
            + "\n- ".join(missing)
        )


def as_float(value: Any, default: float = math.nan) -> float:
    try:
        result = float(value)
        return result if math.isfinite(result) else default
    except (TypeError, ValueError):
        return default


def as_int(value: Any, default: int = 0) -> int:
    try:
        return int(float(value))
    except (TypeError, ValueError):
        return default


def fmt(value: Any, digits: int = 1, suffix: str = "") -> str:
    number = as_float(value)
    return "—" if not math.isfinite(number) else f"{number:.{digits}f}{suffix}"


def pct(value: Any, digits: int = 0) -> str:
    number = as_float(value)
    return "—" if not math.isfinite(number) else f"{100.0 * number:.{digits}f}%"


def pp(value: Any, digits: int = 0) -> str:
    return fmt(value, digits, " pp")


def safe_text(value: Any) -> str:
    if value is None or (isinstance(value, float) and math.isnan(value)):
        return "—"
    return str(value)


@dataclass
class Evidence:
    causal_manifest: dict[str, Any]
    classification: dict[str, Any]
    local_manifest: dict[str, Any]
    propagation_audit: dict[str, Any]
    matched_c00_gate: dict[str, Any]
    atlas_summary: dict[str, Any]
    old_synthesis: str
    cards: pd.DataFrame
    pairs: pd.DataFrame
    condition_summary: pd.DataFrame
    failures: pd.DataFrame
    jacobian: pd.DataFrame
    physics: pd.DataFrame
    event_contact: pd.DataFrame
    propagation: pd.DataFrame
    stage_timing: pd.DataFrame
    chain_order: pd.DataFrame
    old_conditions: pd.DataFrame
    old_joint_profile: pd.DataFrame
    atlas_joint_summary: pd.DataFrame
    source_paths: list[Path]


def validate_and_load() -> Evidence:
    """Validate completion gates before reading any derived CSV."""
    causal_manifest_path = CAUSAL_DIR / "analysis_manifest.json"
    classification_path = CAUSAL_DIR / "classification.json"
    local_manifest_path = LOCAL_DIR / "ANALYSIS_MANIFEST.json"
    local_validation_path = LOCAL_DIR / "JACOBIAN_VALIDATION_PASS.json"
    propagation_audit_path = PROPAGATION_DIR / "ANALYSIS_AUDIT.json"
    matched_gate_path = MATCHED_C00_ROOT / "MATCHED_C00_COMPLETE.json"
    old_seal_path = OLD_MECHANISM_ROOT / "FINAL_EVIDENCE_SEAL.json"
    old_validation_path = OLD_MECHANISM_ROOT / "VALIDATION_PASS.json"
    old_synthesis_path = OLD_MECHANISM_ROOT / "FINAL_MECHANISM_SYNTHESIS.md"
    atlas_complete_path = ATLAS_ROOT / "ANALYSIS_COMPLETE.json"
    matched_result_paths = [
        MATCHED_C00_ROOT / "results" / f"seed{seed}" / "C00.json"
        for seed in range(9201, 9206)
    ]
    gates = [
        causal_manifest_path,
        classification_path,
        local_manifest_path,
        local_validation_path,
        propagation_audit_path,
        matched_gate_path,
        old_seal_path,
        old_validation_path,
        old_synthesis_path,
        atlas_complete_path,
        *matched_result_paths,
    ]
    require_files(gates)
    causal_manifest = load_json(causal_manifest_path)
    classification = load_json(classification_path)
    local_manifest = load_json(local_manifest_path)
    propagation_audit = load_json(propagation_audit_path)
    matched_c00_gate = load_json(matched_gate_path)
    old_seal = load_json(old_seal_path)
    old_validation = load_json(old_validation_path)
    atlas_complete = load_json(atlas_complete_path)
    failures: list[str] = []
    causal_validation = classification.get("validation", {})
    if causal_validation.get("passed") is not True:
        failures.append("analysis_causal_completion validation.passed != true")
    local_validation = local_manifest.get("validation", {})
    if local_validation.get("passed") is not True:
        failures.append("analysis_local_actor validation.passed != true")
    if propagation_audit.get("status") != "complete":
        failures.append("analysis_physical_propagation status != complete")
    if propagation_audit.get("inputs_unchanged") is not True:
        failures.append("analysis_physical_propagation inputs_unchanged != true")
    matched_seed_results = matched_c00_gate.get("seed_results", [])
    matched_successes = [
        as_int(row.get("success_episodes"))
        for row in matched_seed_results
        if isinstance(row, dict)
    ]
    if (
        matched_c00_gate.get("schema")
        != "obs2_v2_1_k_causal_completion/matched_c00_complete/v1"
        or matched_c00_gate.get("status") != "complete"
        or matched_c00_gate.get("training_seeds") != list(range(9201, 9206))
        or matched_c00_gate.get("evaluation_seeds") != list(range(20264401, 20264421))
        or as_int(matched_c00_gate.get("episodes_per_training_seed")) != 20
        or as_int(matched_c00_gate.get("total_episodes")) != 100
        or as_int(matched_c00_gate.get("total_successes")) != 46
        or matched_successes != [18, 0, 9, 0, 19]
        or matched_c00_gate.get("all_checkpoint_and_policy_immutability_gates_passed")
        is not True
    ):
        failures.append("matched C00 completion gate identity/count/immutability is invalid")
    matched_rows_by_seed = {
        as_int(row.get("training_seed")): row
        for row in matched_seed_results
        if isinstance(row, dict)
    }
    for seed, result_path in zip(range(9201, 9206), matched_result_paths):
        row = matched_rows_by_seed.get(seed, {})
        if row.get("result_sha256") != sha256(result_path):
            failures.append(f"matched C00 result hash mismatch: seed {seed}")
    if (
        old_seal.get("schema") != "obs2_v2_1_k_final_evidence_seal/v1"
        or not isinstance(old_seal.get("files"), dict)
        or as_int(old_seal.get("inventory", {}).get("condition_count")) != 59
    ):
        failures.append("legacy mechanism evidence seal schema/inventory is invalid")
    if old_validation.get("passed") is not True:
        failures.append("legacy mechanism validation did not pass")
    if (
        atlas_complete.get("schema") != "formal10_initial3_k1k2_complete/v1"
        or atlas_complete.get("integrity", {}).get("passed") is not True
    ):
        failures.append("atlas analysis completion marker/integrity is invalid")
    if failures:
        raise RuntimeError("Final-report input gate failed:\n- " + "\n- ".join(failures))

    required = [
        CAUSAL_DIR / "per_k_mechanism_cards.csv",
        CAUSAL_DIR / "joint_pair_effects.csv",
        CAUSAL_DIR / "condition_summary.csv",
        CAUSAL_DIR / "failure_decomposition.csv",
        LOCAL_DIR / "jacobian_local_summary_by_seed_event.csv",
        LOCAL_DIR / "shapley_torque_power_by_seed_event.csv",
        LOCAL_DIR / "event_bin_contact_summary.csv",
        PROPAGATION_DIR / "per_channel_propagation.csv",
        PROPAGATION_DIR / "stage_timing_summary.csv",
        PROPAGATION_DIR / "contact_rotation_first_change.csv",
        OLD_MECHANISM_ROOT / "analysis" / "condition_summary.csv",
        OLD_MECHANISM_ROOT / "analysis" / "baseline_joint_profile.csv",
        ATLAS_ROOT / "data" / "joint_terminal_summary.csv",
        ATLAS_ROOT / "data" / "analysis_summary.json",
    ]
    require_files(required)
    source_paths = gates + required
    with old_synthesis_path.open("r", encoding="utf-8-sig") as handle:
        old_synthesis = handle.read()
    return Evidence(
        causal_manifest=causal_manifest,
        classification=classification,
        local_manifest=local_manifest,
        propagation_audit=propagation_audit,
        matched_c00_gate=matched_c00_gate,
        atlas_summary=load_json(ATLAS_ROOT / "data" / "analysis_summary.json"),
        old_synthesis=old_synthesis,
        cards=read_csv(CAUSAL_DIR / "per_k_mechanism_cards.csv"),
        pairs=read_csv(CAUSAL_DIR / "joint_pair_effects.csv"),
        condition_summary=read_csv(CAUSAL_DIR / "condition_summary.csv"),
        failures=read_csv(CAUSAL_DIR / "failure_decomposition.csv"),
        jacobian=read_csv(LOCAL_DIR / "jacobian_local_summary_by_seed_event.csv"),
        physics=read_csv(LOCAL_DIR / "shapley_torque_power_by_seed_event.csv"),
        event_contact=read_csv(LOCAL_DIR / "event_bin_contact_summary.csv"),
        propagation=read_csv(PROPAGATION_DIR / "per_channel_propagation.csv"),
        stage_timing=read_csv(PROPAGATION_DIR / "stage_timing_summary.csv"),
        chain_order=read_csv(PROPAGATION_DIR / "contact_rotation_first_change.csv"),
        old_conditions=read_csv(OLD_MECHANISM_ROOT / "analysis" / "condition_summary.csv"),
        old_joint_profile=read_csv(OLD_MECHANISM_ROOT / "analysis" / "baseline_joint_profile.csv"),
        atlas_joint_summary=read_csv(ATLAS_ROOT / "data" / "joint_terminal_summary.csv"),
        source_paths=source_paths,
    )


def compatibility_audit(evidence: Evidence) -> dict[str, Any]:
    """Audit final analysis schemas and the report's semantic assumptions."""
    required_columns = {
        "per_k_mechanism_cards": (
            evidence.cards,
            {
                "joint",
                "channel",
                "formal_classifications",
                "zero_success_rate_difference_percentage_points",
                "scale_0p5_success_rate",
                "scale_1p5_success_rate",
                "sign_flip_success_rate_difference_percentage_points",
                "static_mean_success_rate_difference_percentage_points",
                "time_permuted_success_rate_difference_percentage_points",
            },
        ),
        "joint_pair_effects": (
            evidence.pairs,
            {
                "joint",
                "sufficiency_reference_uid",
                "sufficiency_reference_pairing",
                "matched_C00_reference_success_rate",
                "sufficiency_success_rate",
                "sufficiency_success_rate_gain_percentage_points",
            },
        ),
        "jacobian_local": (
            evidence.jacobian,
            {
                "training_seed",
                "event_bin",
                "joint",
                "K_channel",
                "observation_channel",
                "positive_fraction",
                "negative_fraction",
                "sign_consistency",
                "derivative_abs_mean",
            },
        ),
        "shapley_physics": (
            evidence.physics,
            {
                "training_seed",
                "event_bin",
                "joint",
                "K_channel",
                "saturation_rate",
                "physical_K_n",
                "shapley_clipped_torque_abs_mean",
                "active_power_proxy_mean",
            },
        ),
        "event_contact": (
            evidence.event_contact,
            {
                "training_seed",
                "event_bin",
                "step_count",
                "saturation_rate_all_joints",
                "clipped_active_torque_abs_mean",
            },
        ),
        "physical_propagation": (
            evidence.propagation,
            {
                "source_joint",
                "source_channel",
                "target_family",
                "target_role",
                "first_sustained_separation_step",
            },
        ),
    }
    missing: list[str] = []
    for name, (frame, columns) in required_columns.items():
        absent = sorted(columns - set(frame.columns))
        if absent:
            missing.append(f"{name}: {absent}")
    if missing:
        raise RuntimeError("Final analysis schema is incompatible:\n- " + "\n- ".join(missing))
    if len(evidence.cards) != 16 or len(evidence.pairs) != 8:
        raise RuntimeError(
            f"Expected 16 per-K cards and 8 joint-pair rows; got {len(evidence.cards)} and {len(evidence.pairs)}"
        )

    pair_reference_ok = bool(
        (evidence.pairs["sufficiency_reference_uid"].astype(str) == "matched::C00").all()
    )
    pair_text_ok = bool(
        evidence.pairs["sufficiency_reference_pairing"]
        .astype(str)
        .str.contains("same training seed", case=False, regex=False)
        .all()
    )
    pair_rate = pd.to_numeric(
        evidence.pairs["matched_C00_reference_success_rate"], errors="coerce"
    )
    pair_rate_ok = bool((pair_rate.sub(0.46).abs() <= 1e-12).all())
    if not (pair_reference_ok and pair_text_ok and pair_rate_ok):
        raise RuntimeError(
            "Joint-pair sufficiency is not consistently bound to matched::C00 at 46/100"
        )

    positive = pd.to_numeric(evidence.jacobian["positive_fraction"], errors="coerce")
    negative = pd.to_numeric(evidence.jacobian["negative_fraction"], errors="coerce")
    reported_consistency = pd.to_numeric(
        evidence.jacobian["sign_consistency"], errors="coerce"
    )
    expected_consistency = pd.concat([positive, negative], axis=1).max(axis=1)
    sign_error = float((reported_consistency - expected_consistency).abs().max())
    if not math.isfinite(sign_error) or sign_error > 1e-12:
        raise RuntimeError(
            "sign_consistency is not max(positive_fraction, negative_fraction)"
        )

    physics = evidence.physics.copy()
    physics["saturation_rate"] = pd.to_numeric(
        physics["saturation_rate"], errors="coerce"
    )
    pivot = physics.pivot_table(
        index=["training_seed", "event_bin", "joint"],
        columns="K_channel",
        values="saturation_rate",
        aggfunc="first",
    )
    if not {"K1", "K2"}.issubset(pivot.columns):
        raise RuntimeError("Per-joint saturation audit needs both K1 and K2 rows")
    channel_saturation_error = float((pivot["K1"] - pivot["K2"]).abs().max())
    if not math.isfinite(channel_saturation_error) or channel_saturation_error > 1e-12:
        raise RuntimeError(
            "K1/K2 rows disagree on joint-total torque saturation rate"
        )

    selected = physics[physics["K_channel"] == "K1"].copy()
    selected["physical_K_n"] = pd.to_numeric(selected["physical_K_n"], errors="coerce")
    selected["weighted_saturation"] = (
        selected["saturation_rate"] * selected["physical_K_n"]
    )
    reconstructed = (
        selected.groupby(["training_seed", "event_bin"], as_index=False)
        .agg(
            weighted_saturation=("weighted_saturation", "sum"),
            weight=("physical_K_n", "sum"),
        )
    )
    reconstructed["expected_all_joints"] = (
        reconstructed["weighted_saturation"] / reconstructed["weight"]
    )
    event_contact = evidence.event_contact.copy()
    event_contact["saturation_rate_all_joints"] = pd.to_numeric(
        event_contact["saturation_rate_all_joints"], errors="coerce"
    )
    saturation_check = reconstructed.merge(
        event_contact[
            ["training_seed", "event_bin", "saturation_rate_all_joints"]
        ],
        on=["training_seed", "event_bin"],
        how="inner",
        validate="one_to_one",
    )
    if len(saturation_check) != len(event_contact):
        raise RuntimeError("All-joints saturation rows do not pair with physics rows")
    all_joint_saturation_error = float(
        (
            saturation_check["expected_all_joints"]
            - saturation_check["saturation_rate_all_joints"]
        )
        .abs()
        .max()
    )
    if (
        not math.isfinite(all_joint_saturation_error)
        or all_joint_saturation_error > 1e-12
    ):
        raise RuntimeError(
            "saturation_rate_all_joints does not equal the weighted joint-step fraction"
        )

    event_counts = (
        evidence.jacobian.groupby("event_bin").size().astype(int).to_dict()
    )
    return {
        "required_columns_passed": True,
        "prelaunch_row_count": int(event_counts.get("official_prelaunch", 0)),
        "events_present": sorted(str(item) for item in event_counts),
        "sign_consistency_definition_verified": True,
        "sign_consistency_max_abs_error": sign_error,
        "joint_total_saturation_K1_K2_max_abs_error": channel_saturation_error,
        "all_joints_saturation_reconstruction_max_abs_error": all_joint_saturation_error,
        "matched_C00_reference_verified": True,
        "matched_C00_success_rate": 0.46,
        "matched_C00_pairing": "same training seed and evaluation seeds 20264401..20264420",
    }


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_run_font(
    run,
    *,
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = PRESET["body_font"]
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), PRESET["body_font"])
    rfonts.set(qn("w:hAnsi"), PRESET["body_font"])
    rfonts.set(qn("w:eastAsia"), PRESET["east_asia_font"])
    rfonts.set(qn("w:cs"), PRESET["body_font"])
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, size: float, color: str = "17212B", bold: bool = False) -> None:
    style.font.name = PRESET["body_font"]
    style.font.size = Pt(size)
    style.font.color.rgb = rgb(color)
    style.font.bold = bold
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), PRESET["body_font"])
    rfonts.set(qn("w:hAnsi"), PRESET["body_font"])
    rfonts.set(qn("w:eastAsia"), PRESET["east_asia_font"])
    rfonts.set(qn("w:cs"), PRESET["body_font"])


def set_spacing_xml(paragraph_format, before_pt: float, after_pt: float, line_twips: int) -> None:
    paragraph_format.space_before = Pt(before_pt)
    paragraph_format.space_after = Pt(after_pt)
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    ppr = paragraph_format._element.get_or_add_pPr()
    spacing = ppr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        ppr.append(spacing)
    spacing.set(qn("w:before"), str(round(before_pt * 20)))
    spacing.set(qn("w:after"), str(round(after_pt * 20)))
    spacing.set(qn("w:line"), str(line_twips))
    spacing.set(qn("w:lineRule"), "auto")


def configure_document_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(PRESET["page_width_in"])
    section.page_height = Inches(PRESET["page_height_in"])
    section.top_margin = Inches(PRESET["margin_in"])
    section.bottom_margin = Inches(PRESET["margin_in"])
    section.left_margin = Inches(PRESET["margin_in"])
    section.right_margin = Inches(PRESET["margin_in"])
    section.header_distance = Inches(PRESET["header_footer_distance_in"])
    section.footer_distance = Inches(PRESET["header_footer_distance_in"])

    normal = doc.styles["Normal"]
    set_style_font(normal, PRESET["body_size_pt"], PRESET["ink"])
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_spacing_xml(normal.paragraph_format, 0, PRESET["body_after_pt"], PRESET["body_line_twips"])
    normal.paragraph_format.widow_control = True

    for name, size, color, before, after in (
        ("Heading 1", PRESET["h1_size_pt"], PRESET["h1_color"], PRESET["h1_before_pt"], PRESET["h1_after_pt"]),
        ("Heading 2", PRESET["h2_size_pt"], PRESET["h2_color"], PRESET["h2_before_pt"], PRESET["h2_after_pt"]),
        ("Heading 3", PRESET["h3_size_pt"], PRESET["h3_color"], PRESET["h3_before_pt"], PRESET["h3_after_pt"]),
    ):
        style = doc.styles[name]
        set_style_font(style, size, color, bold=True)
        set_spacing_xml(style.paragraph_format, before, after, 240)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    if "Report Caption" not in doc.styles:
        style = doc.styles.add_style("Report Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        style = doc.styles["Report Caption"]
    set_style_font(style, 9.0, PRESET["muted"])
    style.font.italic = True
    set_spacing_xml(style.paragraph_format, 3, 8, 240)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style.paragraph_format.keep_with_next = False

    if "Table Source" not in doc.styles:
        style = doc.styles.add_style("Table Source", WD_STYLE_TYPE.PARAGRAPH)
    else:
        style = doc.styles["Table Source"]
    set_style_font(style, 8.5, PRESET["muted"])
    style.font.italic = True
    set_spacing_xml(style.paragraph_format, 4, 4, 240)


def set_paragraph_keep(paragraph, *, next_: bool = False, together: bool = False) -> None:
    paragraph.paragraph_format.keep_with_next = next_
    paragraph.paragraph_format.keep_together = together
    paragraph.paragraph_format.widow_control = True


def paragraph_shading(paragraph, fill: str, left_border: str | None = None) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:fill"), fill)
    if left_border:
        borders = ppr.find(qn("w:pBdr"))
        if borders is None:
            borders = OxmlElement("w:pBdr")
            ppr.append(borders)
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), left_border)
        borders.append(left)


def add_callout(doc: Document, label: str, text: str, *, kind: str = "info"):
    palette = {
        "info": (PRESET["light_blue"], PRESET["blue"]),
        "evidence": (PRESET["light_gold"], PRESET["gold"]),
        "risk": (PRESET["light_red"], PRESET["red"]),
    }
    fill, border = palette[kind]
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.14)
    paragraph.paragraph_format.right_indent = Inches(0.05)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.paragraph_format.line_spacing = 1.20
    paragraph_shading(paragraph, fill, border)
    label_run = paragraph.add_run(f"{label}  ")
    set_run_font(label_run, size=10.5, color=border, bold=True)
    text_run = paragraph.add_run(text)
    set_run_font(text_run, size=10.5, color=PRESET["ink"])
    set_paragraph_keep(paragraph, together=True)
    return paragraph


def add_field(paragraph, instruction: str, result: str = "") -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = result
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def configure_section(section, *, header_text: str | None = None, first_page: bool = False) -> None:
    section.page_width = Inches(PRESET["page_width_in"])
    section.page_height = Inches(PRESET["page_height_in"])
    section.top_margin = Inches(PRESET["margin_in"])
    section.bottom_margin = Inches(PRESET["margin_in"])
    section.left_margin = Inches(PRESET["margin_in"])
    section.right_margin = Inches(PRESET["margin_in"])
    section.header_distance = Inches(PRESET["header_footer_distance_in"])
    section.footer_distance = Inches(PRESET["header_footer_distance_in"])
    section.different_first_page_header_footer = first_page
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    if header_text:
        hp = section.header.paragraphs[0]
        hp.text = ""
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        hp.paragraph_format.space_after = Pt(2)
        run = hp.add_run(header_text)
        set_run_font(run, size=8.5, color=PRESET["muted"], bold=True)
        ppr = hp._p.get_or_add_pPr()
        borders = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "5")
        bottom.set(qn("w:color"), "D9E2E8")
        borders.append(bottom)
        ppr.append(borders)
        fp = section.footer.paragraphs[0]
        fp.text = ""
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        prefix = fp.add_run("Page ")
        set_run_font(prefix, size=8.5, color=PRESET["muted"])
        add_field(fp, "PAGE", "1")
        suffix = fp.add_run("")
        set_run_font(suffix, size=8.5, color=PRESET["muted"])


def add_cover(doc: Document) -> None:
    section = doc.sections[0]
    configure_section(section, first_page=True)
    section.first_page_header.paragraphs[0].text = ""
    section.first_page_footer.paragraphs[0].text = ""
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(92)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    run = kicker.add_run("Per-Joint K1/K2 Causal-Mechanism Study")
    set_run_font(run, size=10.5, color=PRESET["gold"], bold=True)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)
    run = title.add_run("From Local Observations to Forward Rolling")
    set_run_font(run, size=30, color=PRESET["navy"], bold=True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(8)
    run = subtitle.add_run("Policy Response, Torque and Power, Contact Propagation, and Frozen Causal Intervention Across 16 K Channels")
    set_run_font(run, size=15, color=PRESET["dark_blue"], bold=False)
    contract = doc.add_paragraph()
    contract.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contract.paragraph_format.space_after = Pt(72)
    run = contract.add_run(
        "Observations permanently locked: [delta_theta, theta_dot] per joint - Actions permanently locked: [K1, K2] per joint\n"
        "Formal seeds 9201-9205 - checkpoint 1500 - deterministic frozen CPU evaluation"
    )
    set_run_font(run, size=10.5, color=PRESET["muted"])
    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date.paragraph_format.space_after = Pt(6)
    run = date.add_run("Integrated Research Report - 2026-08-04")
    set_run_font(run, size=12, color=PRESET["navy"], bold=True)
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run("Read-only synthesis of frozen evidence; no retraining, checkpoint modification, or observation/action-channel change")
    set_run_font(run, size=9.5, color=PRESET["muted"], italic=True)


def add_body_section(doc: Document) -> None:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(
        section,
        header_text="Per-Joint K1/K2 Causal-Mechanism Completion - obs2_roll_repro_v2.1",
        first_page=False,
    )


def create_numbering(doc: Document, *, ordered: bool) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = (max(abstract_ids) + 1) if abstract_ids else 1
    num_id = (max(num_ids) + 1) if num_ids else 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if ordered else "bullet")
    level.append(num_fmt)
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), "%1." if ordered else "•")
    level.append(text)
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    level.append(justification)
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), str(PRESET["list_text_dxa"]))
    tabs.append(tab)
    ppr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(PRESET["list_text_dxa"]))
    ind.set(qn("w:hanging"), str(PRESET["list_hanging_dxa"]))
    ppr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), str(PRESET["list_after_twips"]))
    spacing.set(qn("w:line"), str(PRESET["list_line_twips"]))
    spacing.set(qn("w:lineRule"), "auto")
    ppr.append(spacing)
    level.append(ppr)
    rpr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), PRESET["body_font"])
    fonts.set(qn("w:hAnsi"), PRESET["body_font"])
    fonts.set(qn("w:eastAsia"), PRESET["east_asia_font"])
    rpr.append(fonts)
    level.append(rpr)
    abstract.append(level)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_list_item(doc: Document, text: str, num_id: int, *, bold_lead: str | None = None):
    paragraph = doc.add_paragraph()
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    ppr.append(num_pr)
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_run_font(lead, bold=True)
        run = paragraph.add_run(text[len(bold_lead) :])
        set_run_font(run)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)
    set_paragraph_keep(paragraph, together=True)
    return paragraph


def set_cell_margins(cell, margins: Mapping[str, int] | None = None) -> None:
    margins = margins or PRESET["cell_margins_dxa"]
    tc = cell._tc
    tcpr = tc.get_or_add_tcPr()
    tc_mar = tcpr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tcpr.append(tc_mar)
    for key in ("top", "start", "bottom", "end"):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margins[key]))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: Sequence[int]) -> None:
    if sum(widths_dxa) != PRESET["usable_width_dxa"]:
        raise ValueError(f"Table widths must sum to 9360 DXA: {widths_dxa}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(PRESET["usable_width_dxa"]))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(PRESET["table_indent_dxa"]))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        row.height = None
        row._tr.get_or_add_trPr()
        for index, cell in enumerate(row.cells):
            tcpr = cell._tc.get_or_add_tcPr()
            tcw = tcpr.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                tcpr.append(tcw)
            tcw.set(qn("w:w"), str(widths_dxa[index]))
            tcw.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def shade_cell(cell, fill: str) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def repeat_table_header(row) -> None:
    trpr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    trpr.append(header)


def add_data_table(
    doc: Document,
    headers: Sequence[str],
    rows: Sequence[Sequence[Any]],
    widths_dxa: Sequence[int],
    *,
    font_size: float = 8.7,
    alignments: Sequence[int] | None = None,
) -> Any:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header = table.rows[0]
    repeat_table_header(header)
    for index, value in enumerate(headers):
        cell = header.cells[index]
        shade_cell(cell, PRESET["table_header_fill"])
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(1)
        paragraph.paragraph_format.space_after = Pt(1)
        run = paragraph.add_run(str(value))
        set_run_font(run, size=font_size, color=PRESET["navy"], bold=True)
    for row_values in rows:
        row = table.add_row()
        for index, value in enumerate(row_values):
            cell = row.cells[index]
            paragraph = cell.paragraphs[0]
            paragraph.alignment = (
                alignments[index] if alignments else WD_ALIGN_PARAGRAPH.LEFT
            )
            paragraph.paragraph_format.space_before = Pt(1)
            paragraph.paragraph_format.space_after = Pt(1)
            paragraph.paragraph_format.line_spacing = 1.10
            run = paragraph.add_run(safe_text(value))
            set_run_font(run, size=font_size, color=PRESET["ink"])
    set_table_geometry(table, widths_dxa)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def set_alt_text(inline_shape, title: str, description: str) -> None:
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", description)


def add_figure(
    doc: Document,
    path: Path,
    figure_id: str,
    title: str,
    interpretation: str,
    *,
    max_width_in: float = 6.20,
    max_height_in: float = 6.15,
) -> None:
    if not path.is_file():
        raise FileNotFoundError(path)
    with Image.open(path) as image:
        ratio = image.width / max(image.height, 1)
    width = min(max_width_in, max_height_in * ratio)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(2)
    set_paragraph_keep(paragraph, next_=True, together=True)
    shape = paragraph.add_run().add_picture(str(path), width=Inches(width))
    set_alt_text(shape, f"{figure_id} {title}", interpretation)
    caption = doc.add_paragraph(style="Report Caption")
    run = caption.add_run(f"{figure_id} | {title}. {interpretation}")
    set_run_font(run, size=9.0, color=PRESET["muted"], italic=True)


def add_source_note(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="Table Source")
    run = paragraph.add_run(text)
    set_run_font(run, size=8.5, color=PRESET["muted"], italic=True)


def row_for(frame: pd.DataFrame, **filters: Any) -> pd.Series | None:
    subset = frame
    for key, value in filters.items():
        if key not in subset.columns:
            return None
        subset = subset[subset[key] == value]
    return None if subset.empty else subset.iloc[0]


def condition_rate(evidence: Evidence, condition_id: str) -> float:
    row = row_for(evidence.old_conditions, condition_id=condition_id)
    return math.nan if row is None else as_float(row.get("mean_success_rate"))


def card_row(evidence: Evidence, joint: str, channel: str) -> pd.Series:
    row = row_for(evidence.cards, joint=joint, channel=channel)
    if row is None:
        raise RuntimeError(f"Missing per-K card: {joint}-{channel}")
    return row


def class_labels(row: pd.Series) -> list[str]:
    raw = str(row.get("formal_classifications", "no_frozen_threshold_label"))
    labels = [CLASS_CN.get(item, item) for item in raw.split(";") if item]
    return labels or [CLASS_CN["no_frozen_threshold_label"]]


def baseline_profile(evidence: Evidence, joint: str, channel: str) -> dict[str, float]:
    row = row_for(evidence.old_joint_profile, joint=joint)
    if row is None:
        return {"mean": math.nan, "abs_mean": math.nan, "positive": math.nan}
    prefix = channel
    return {
        "mean": as_float(row.get(f"{prefix}_mean")),
        "abs_mean": as_float(row.get(f"{prefix}_abs_mean")),
        "positive": as_float(row.get(f"{prefix}_positive_fraction")),
    }


def dominant_jacobian(evidence: Evidence, joint: str, channel: str) -> dict[str, Any]:
    subset = evidence.jacobian[
        (evidence.jacobian["joint"] == joint)
        & (evidence.jacobian["K_channel"] == channel)
    ].copy()
    if subset.empty:
        return {"input": "—", "abs": math.nan, "event": "—", "sign": math.nan}
    subset["abs_derivative"] = pd.to_numeric(
        subset["derivative_abs_mean"], errors="coerce"
    )
    grouped = subset.groupby("observation_channel", as_index=False)["abs_derivative"].median()
    input_name = str(grouped.sort_values("abs_derivative", ascending=False).iloc[0]["observation_channel"])
    # sign_consistency is defined by the analyzer as max(positive_fraction,
    # negative_fraction) within one seed/event bin.  Aggregate it across seeds
    # by the median; never reinterpret it as cross-seed sign agreement.
    selected = (
        subset[subset["observation_channel"] == input_name]
        .groupby("event_bin", as_index=False)
        .agg(
            abs_derivative=("abs_derivative", "median"),
            positive_fraction=("positive_fraction", "median"),
            negative_fraction=("negative_fraction", "median"),
            sign_consistency=("sign_consistency", "median"),
        )
        .sort_values("abs_derivative", ascending=False)
    )
    best = selected.iloc[0]
    return {
        "input": "Δθ" if input_name == "delta_theta" else "θ̇",
        "abs": as_float(best.get("derivative_abs_mean")),
        "event": EVENT_CN.get(str(best.get("event_bin")), str(best.get("event_bin"))),
        "sign": as_float(best.get("positive_fraction")) - as_float(best.get("negative_fraction")),
        "consistency": as_float(best.get("sign_consistency")),
    }


def physics_summary(evidence: Evidence, joint: str, channel: str) -> dict[str, Any]:
    subset = evidence.physics[
        (evidence.physics["joint"] == joint)
        & (evidence.physics["K_channel"] == channel)
    ].copy()
    if subset.empty:
        return {"torque": math.nan, "power": math.nan, "saturation": math.nan, "event": "—"}
    grouped = (
        subset.groupby("event_bin", as_index=False)
        .agg(
            torque=("shapley_clipped_torque_abs_mean", "median"),
            power=("active_power_proxy_mean", "median"),
            saturation=("saturation_rate", "median"),
        )
        .sort_values("torque", ascending=False)
    )
    best = grouped.iloc[0]
    return {
        "torque": as_float(best["torque"]),
        "power": as_float(best["power"]),
        "saturation": as_float(best["saturation"]),
        "event": EVENT_CN.get(str(best["event_bin"]), str(best["event_bin"])),
    }


def propagation_summary(evidence: Evidence, joint: str, channel: str) -> dict[str, Any]:
    subset = evidence.propagation[
        (evidence.propagation["source_joint"] == joint)
        & (evidence.propagation["source_channel"] == channel)
        & (evidence.propagation["target_role"].isin(["neighbor", "far"]))
        & (evidence.propagation["target_family"].isin(["observation", "position", "K", "tau"]))
    ].copy()
    if subset.empty:
        return {"step": math.nan, "family": "—", "role": "—", "detected": 0.0}
    subset["first"] = pd.to_numeric(
        subset["first_sustained_separation_step"], errors="coerce"
    )
    detected = subset.dropna(subset=["first"])
    if detected.empty:
        return {"step": math.nan, "family": "—", "role": "—", "detected": 0.0}
    grouped = (
        detected.groupby(["target_family", "target_role"], as_index=False)["first"]
        .median()
        .sort_values("first")
    )
    best = grouped.iloc[0]
    return {
        "step": as_float(best["first"]),
        "family": str(best["target_family"]),
        "role": str(best["target_role"]),
        "detected": len(detected) / len(subset),
    }


def atlas_success(evidence: Evidence) -> tuple[list[int], list[int]]:
    mapping = evidence.atlas_summary.get("formal_terminal_success_common_metric", {})
    r0 = [as_int(mapping.get(f"formal_{seed}_R0")) for seed in range(9201, 9206)]
    rroll = [as_int(mapping.get(f"formal_{seed}_Rroll")) for seed in range(9201, 9206)]
    return r0, rroll


def find_figure(base: Path, *relative_candidates: str) -> Path | None:
    for candidate in relative_candidates:
        path = base / candidate
        if path.is_file():
            return path
    return None


def maybe_add_figure(
    doc: Document,
    path: Path | None,
    figure_id: str,
    title: str,
    interpretation: str,
    **kwargs: Any,
) -> None:
    if path is not None:
        add_figure(doc, path, figure_id, title, interpretation, **kwargs)


def add_toc(doc: Document) -> None:
    doc.add_heading("Contents", level=1)
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(12)
    add_field(paragraph, 'TOC \\o "1-3" \\h \\z \\u', "Update this table of contents in Word")
    add_callout(
        doc,
        "Suggested reading order",
        "Read the executive summary and validated-conclusions list first, then the 16 K mechanism cards. Jacobian, Shapley, and propagation sections explain why; frozen-intervention sections adjudicate necessity and sufficiency.",
        kind="info",
    )


def add_executive_summary(doc: Document, evidence: Evidence) -> None:
    doc.add_page_break()
    doc.add_heading("Executive Summary", level=1)
    r0, rroll = atlas_success(evidence)
    c00, c01, c10, c11 = [condition_rate(evidence, item) for item in ("C00", "C01", "C10", "C11")]
    add_callout(
        doc,
        "Core answer",
        "Forward rolling is not caused by one fixed K value. It emerges jointly from 16 K channels produced by eight local closed-loop policies at their own [delta_theta, theta_dot] states. The direct policy mapping contains eight local 2x2 blocks; cross-joint cooperation arises from closed-loop body dynamics, contact, and local observations rather than an actor directly reading other joints.",
        kind="evidence",
    )
    doc.add_paragraph(
        f"In formal frozen evaluation, the five R0 training seeds produced {r0}/20 successful episodes, {sum(r0)}/100 total; "
        f"Rroll produced {rroll}/20, {sum(rroll)}/100 total. Rroll increased the number of seeds meeting threshold from "
        f"{sum(value >= 10 for value in r0)}/5 to {sum(value >= 10 for value in rroll)}/5. "
        "The evidence therefore supports the conclusion that the reward stabilizes occasional rolling into cross-seed reproducible rolling."
    )
    doc.add_paragraph(
        f"Legacy frozen channel-factor success rates were C00={pct(c00)}, C01={pct(c01)}, "
        f"C10={pct(c10)}, and C11={pct(c11)}. K1 and K2 are not two independently additive knobs; "
        f"their joint interaction is approximately {fmt(c11 - c10 - c01 + c00, 2)}."
    )
    doc.add_paragraph(
        "To remove cross-initial-state ambiguity in whole-joint sufficiency, this study adds matched C00. "
        "Using the same training seeds, Rroll environment, and evaluation seeds 20264401-20264420 as the 113-condition matrix, "
        "the five seeds yielded [18,0,9,0,19]/20 successful episodes, 46/100 total. All later whole-joint K1+K2 sufficiency uses this same-initial-state C00 reference."
    )
    doc.add_paragraph(
        "This completion study advances beyond what K the policy might output to four evidence layers: "
        "the local Jacobian measures each K's response to its two local observations; Shapley decomposition assigns active torque and power proxy after clipping; "
        "propagation analysis tracks how local changes travel through contact and body; frozen zeroing, flipping, dose, static replacement, and temporal permutation adjudicate necessity and timing dependence."
    )
    add_data_table(
        doc,
        ["Evidence layer", "Question answered", "Independent inference unit", "Role in this report"],
        [
            ["Formal endpoint evaluation", "Whether forward rolling is stable", "Five training seeds", "Primary feasibility conclusion"],
            ["Frozen intervention", "Whether a K is necessary, sufficient, or timing-critical", "Five training seeds; 20 paired initial states per seed", "Primary causal conclusion"],
            ["Local Jacobian", "How observation changes alter local K", "Analytic derivatives at visited states", "Policy-mechanism interpretation"],
            ["Shapley torque/power", "How K1/K2 pass through clipping", "Visited trajectories and event phases", "Control-boundary interpretation"],
            ["Physical propagation", "When a local intervention affects adjacent joints, contact, and rotation", "One paired trajectory per condition per seed", "Propagation-order diagnosis"],
            ["Historical 101x101 atlas", "What the policy may output on a two-dimensional local-state surface", "195 checkpoints", "Descriptive background, not a causal criterion"],
        ],
        [1250, 2740, 2470, 2900],
        font_size=8.6,
    )
    add_source_note(doc, "Table note: 100 evaluation episodes are not 100 independent policies; scientific inference uses five training seeds.")


def add_contract_and_architecture(doc: Document, evidence: Evidence) -> None:
    doc.add_page_break()
    doc.add_heading("1 Research Contract and Key Architectural Correction", level=1)
    doc.add_heading("1.1 Permanently Locked Variables", level=2)
    bullets = create_numbering(doc, ordered=False)
    for text in (
        "Observations remain [delta_theta, theta_dot] per joint; contact, phase, body pose, and whole-body state are not added.",
        "Actions remain [K1, K2] per joint. Deterministic actions use policy loc and map to physical gains as K=100*loc.",
        "Active control torque is clip(K1*delta_theta + K2*theta_dot, -9, +9); this analysis changes neither physics, PPO, formal checkpoints, nor the five-part success threshold.",
        "The report reads only frozen checkpoint-1500 results and derived analyses; it does not continue training or select intermediate checkpoints.",
    ):
        add_list_item(doc, text, bullets)
    doc.add_heading("1.2 Key Architecture Correction: Not a 16-Input Whole-Body Actor", level=2)
    add_callout(
        doc,
        "Required correction",
        "The formal policy contains eight local networks with no parameter sharing. Network j reads only local [delta_theta_j, theta_dot_j] and outputs only local [K1_j, K2_j]. The 16x16 policy Jacobian is therefore structurally block diagonal with eight 2x2 blocks; every direct cross-joint partial derivative is zero.",
        kind="risk",
    )
    doc.add_paragraph(
        "This does not make joints physically independent. delta_theta_j is constructed from adjacent geometry; a local K change first alters local torque and pose, then propagates through the flexible body, ground contact, support transfer, and new local observations. "
        "Direct policy coupling (absent) must be distinguished from closed-loop physical coupling (quantified here)."
    )
    structural = find_figure(LOCAL_DIR, "figures/J01_direct_jacobian_structure.png")
    maybe_add_figure(
        doc,
        structural,
        "Figure 1",
        "Structural mask of the 16x16 policy Jacobian",
        "Eight 2x2 diagonal blocks contain allowed local derivatives; cross-joint blocks are structural zeros. The figure describes only the actor's direct mapping and does not deny later physical propagation.",
    )
    doc.add_heading("1.3 Success Threshold and Evidence Boundary", level=2)
    add_data_table(
        doc,
        ["Joint per-episode criterion", "Threshold", "Meaning"],
        [
            ["Rolling pulses", ">=4", "Not a single fall"],
            ["Target-direction net rotation", ">=360 degrees", "At least one full turn"],
            ["Directional consistency", ">=0.70", "Rotation predominantly follows the target direction"],
            ["Forward distance", ">=1 body length", "Rolling must cause displacement"],
            ["Mean pulse interval", "<=250 steps", "Pulses must be sustained rather than sparse"],
        ],
        [2200, 1200, 5960],
        font_size=9.0,
    )


def add_evidence_inventory(doc: Document, evidence: Evidence) -> None:
    doc.add_page_break()
    doc.add_heading("2 Data, Methods, and Evidence Integrity", level=1)
    doc.add_paragraph(
        "This report uses three new analysis directories and two sealed historical evidence repositories. Every input passes completion gates before generation; SHA-256 is checked again during generation, and delivery is refused if an input changes."
    )
    new_condition_count = evidence.matched_c00_gate.get(
        "parent_condition_count_unchanged", 113
    )
    old_condition_count = evidence.classification.get("validation", {}).get(
        "legacy_metadata_condition_count", "—"
    )
    add_data_table(
        doc,
        ["Asset", "Scale", "Purpose", "Completion gate"],
        [
            ["New causal-completion matrix", f"{new_condition_count} new conditions", "Per-K zeroing/dose/flip/static/permutation and whole-pair necessity/sufficiency", "analysis_manifest.json"],
            ["Matched-C00 control", "5 seeds x 20 shared initial states; 46/100", "Same-seed, same-evaluation-seed paired reference for whole-joint K1+K2 sufficiency", "MATCHED_C00_COMPLETE.json"],
            ["Legacy mechanism matrix", f"{old_condition_count} legacy conditions", "Channel factors, sign space, K2 dose, and legacy single-channel transplant", "FINAL_EVIDENCE_SEAL.json"],
            ["Local-policy analysis", "5 seeds x 8 joints x 2x2 derivatives x event bins", "Direct observation-to-K rule", "JACOBIAN_VALIDATION_PASS.json"],
            ["Physical propagation", f"{evidence.propagation_audit.get('paired_comparison_count', 'n/a')} paired comparisons", "Ordering of K/torque/pose/contact/rotation", "ANALYSIS_AUDIT.json"],
            ["Historical response atlas", "13 runs, 195 checkpoints, 710 figures", "Training evolution and two-dimensional state-surface context", "ANALYSIS_COMPLETE.json"],
        ],
        [1500, 2080, 3460, 2320],
        font_size=8.4,
    )
    doc.add_heading("2.1 Local Jacobian", level=2)
    doc.add_paragraph(
        "For each local network, compute d[K1,K2]/d[delta_theta,theta_dot] analytically and bin actually visited states by official prelaunch, pulse quintile, and rolling-outside-pulse events. Validate independently against autograd and central finite differences; the maximum cross-joint block magnitude must remain numerically zero."
    )
    doc.add_heading("2.2 Clipping-Aware Shapley Decomposition", level=2)
    doc.add_paragraph(
        "Let u1=K1*delta_theta, u2=K2*theta_dot, and f(u)=clip(u,-9,+9). The report uses exact two-channel Shapley allocation: "
        "phi1=0.5[f(u1)-f(0)+f(u1+u2)-f(u2)] and phi2=0.5[f(u2)-f(0)+f(u1+u2)-f(u1)]. "
        "phi1+phi2 exactly reconstructs executed clipped torque; phi*theta_dot is an active-control power proxy, not exact energy over ten physics substeps."
    )
    doc.add_heading("2.3 Frozen Intervention Criteria", level=2)
    doc.add_paragraph(
        "Strong necessity, necessary contribution, timing-critical status, approximate equivalence, and sufficiency all use prefrozen rules for success-rate differences, continuous-metric differences, and cross-seed consistency. "
        "The 20 paired initial states per seed are repeated measurements, not 100 independently learned policies."
    )
    validation = find_figure(LOCAL_DIR, "figures/V01_jacobian_validation.png")
    maybe_add_figure(
        doc,
        validation,
        "Figure 2",
        "Independent numerical validation of the local Jacobian",
        "Analytic derivatives are compared separately with autograd and central finite differences; the error distribution shows that phase derivatives are not plotting approximations.",
        max_height_in=4.6,
    )


def add_project_evidence_timeline(doc: Document, evidence: Evidence) -> None:
    """Summarize the chat-led evolution from a phenotype to a causal study."""
    doc.add_page_break()
    doc.add_heading("3 Project Conversations and Evidence-Evolution Timeline", level=1)
    doc.add_paragraph(
        "The research question did not arrive fully formed; it converged through conversations, training, evaluator repair, and mechanism experiments. "
        "The key transitions below prevent later conclusions from being written as though they were known from the beginning."
    )
    r0, rroll = atlas_success(evidence)
    add_data_table(
        doc,
        ["Stage", "Evidence then available", "Question then answerable", "Why the next step was required"],
        [
            [
                "1. Initial s0 observation",
                "s0 rolled during early one-dimensional-observation/legacy training, then degraded at later checkpoints.",
                "Established existence of a rolling phenotype but not a stable mechanism.",
                "A single-seed, single-period observation may be accidental or training degradation and requires cross-seed testing.",
            ],
            [
                "2. Four seeds and static atlases",
                "Longitudinal four-seed comparison showed only s0 as a rolling donor; two- and one-dimensional response plots visualized K surfaces.",
                "Localized candidate K distributions and training evolution.",
                "Static plots show only what may be output, not necessity, sufficiency, or actual control effect.",
            ],
            [
                "3. Reward-only reproduction v2/v2.1",
                "Keep observations [delta_theta,theta_dot], actions [K1,K2], PPO, and physics unchanged; alter only reward and pair training from batch 0.",
                "Showed rolling can be learned from scratch without expanding channels.",
                "A pilot establishes candidate feasibility only; independent formal seeds and a fixed endpoint are required.",
            ],
            [
                "4. Five formal seeds",
                f"Formal frozen Rroll evaluation was {rroll}/20, {sum(rroll)}/100 total, with 5/5 seeds meeting threshold.",
                "Supported robust cross-seed rolling reproduction.",
                "Same-seed R0 controls remained necessary to separate reward effects from initialization/training randomness.",
            ],
            [
                "5. R0 false-negative correction",
                f"After repairing evaluator fields, corrected R0 was {r0}/20, {sum(r0)}/100 total, rather than never rolling.",
                "Reward should be described as stabilizing occasional rolling rather than creating motion from absolute zero.",
                "R0/Rroll K structure and frozen interventions must be compared, not only success labels.",
            ],
            [
                "6. Legacy 59-condition mechanism study",
                "5 seeds x 59 conditions x 20 episodes; channel factors, K1 sign space, K2 dose/timing, and per-joint replacement.",
                "Established the J01/J02 boundary, K1 sign scaffold, K2 timing, and nonadditive K1/K2 interaction.",
                "Legacy K1 necessity did not use the full C11 background; per-K dose, flip, static replacement, and whole-pair sufficiency remained missing.",
            ],
            [
                "7. Current 113-condition completion",
                "Per-K zero, 0.5x, 1.5x, sign-flip, static-mean, and time-permutation interventions on full C11, plus whole-joint K1+K2 interventions and matched C00 at 46/100.",
                "Supports separate necessity, timing, and equivalence boundaries for 16 K channels and whole-joint sufficiency using the same seeds and initial states 20264401-20264420.",
                "Success rates still require joint interpretation with local policy response and physical propagation to form a mechanism explanation.",
            ],
            [
                "8. Architecture/Jacobian correction",
                "Checkpoint and source audit confirmed an actor containing eight local 2-to-2 networks without shared parameters.",
                "The direct 16x16 Jacobian is strictly block diagonal; direct cross-joint partial derivatives are zero.",
                "Cross-joint cooperation must instead be explained through closed-loop torque, body, contact, and new-observation propagation.",
            ],
        ],
        [1320, 2740, 2440, 2860],
        font_size=7.8,
    )
    add_callout(
        doc,
        "Overall correction produced by the timeline",
        "The conclusion evolved from s0 can roll to Rroll rolls robustly across seeds and local-state K1/K2 responses form whole-body rolling through clipped torque and body-contact feedback. The first is an observation; the second is the mechanism question requested by the adviser.",
        kind="evidence",
    )


def system_conclusions(evidence: Evidence) -> list[str]:
    r0, rroll = atlas_success(evidence)
    c00, c01, c10, c11 = [condition_rate(evidence, item) for item in ("C00", "C01", "C10", "C11")]
    k2_zero = condition_rate(evidence, "K2_SCALE_0")
    k2_half = condition_rate(evidence, "K2_SCALE_0P5")
    k2_one = condition_rate(evidence, "C11")
    k2_one_half = condition_rate(evidence, "K2_SCALE_1P5")
    static = condition_rate(evidence, "K2_CALIBRATION_STATIC_MEAN")
    permuted = condition_rate(evidence, "K2_CALIBRATION_PERMUTED_TEMPLATE")
    results = [
        f"Formal reproducibility: Rroll achieved {sum(rroll)}/100 with 5/5 seeds meeting their threshold; R0 achieved {sum(r0)}/100 with {sum(v >= 10 for v in r0)}/5 seeds meeting threshold. The rolling reward's main value is cross-seed stabilization.",
        "Stable rolling requires no observation/action expansion: all evidence still uses per-joint [delta_theta, theta_dot] -> [K1,K2], and formal training starts at batch 0 without a historical checkpoint.",
        "The architecture contains eight independent local 2-to-2 actors; the direct 16x16 Jacobian must be block diagonal with eight 2x2 blocks. Only body-contact-observation feedback explains cross-joint cooperation.",
        f"K1/K2 interact nonadditively: C00={pct(c00)}, C01={pct(c01)}, C10={pct(c10)}, C11={pct(c11)}, interaction approximately {fmt(c11-c10-c01+c00, 2)}. Optimizing only one channel cannot reproduce the complete mechanism.",
        "The whole-joint sufficiency control is complete: matched C00 is [18,0,9,0,19]/20 (46/100) using the same seeds and evaluation seeds 20264401-20264420, so new whole-pair transplants are strictly paired by initial state. Legacy single-channel sufficiency remains an internal legacy-study comparison.",
        "Legacy frozen interventions support a K1 spatial sign scaffold of positive J01 and negative J02-J08; all-positive, all-negative, alternating, mirrored, and expanded anterior-positive regions cannot replace it.",
        "J01/J02 form the key sign boundary: positive posterior J01 K1 and strongly negative J02 K1 convert local launch into whole-body curling. The posterior is not independently sufficient and requires downstream closed-loop propagation.",
        f"K2 is not a removable decoration: global K2=0 succeeds at {pct(k2_zero)}; 0.5x, 1x, and 1.5x achieve {pct(k2_half)}, {pct(k2_one)}, and {pct(k2_one_half)}, showing an effective amplitude window rather than monotonic benefit.",
        f"K2 timing matters more than a static mean: static-template success is {pct(static)} and time-permuted success {pct(permuted)}. K2 therefore dynamically shapes phase, injection, and dissipation rather than acting as a fixed bias.",
        "Frozen K distributions and the legacy 101x101 atlas jointly reject the claim that larger gain makes rolling easier: aggregate Rroll K1/K2 RMS is below R0, and robustness arises from sign, state dependence, and temporal structure.",
        "A K value is not itself a control effect. Explanation requires aligning K, local observations, clipped torque, power proxy, contact, and body rotation in time.",
    ]
    return results


def channel_conclusion(evidence: Evidence, joint: str, channel: str) -> str:
    card = card_row(evidence, joint, channel)
    profile = baseline_profile(evidence, joint, channel)
    jac = dominant_jacobian(evidence, joint, channel)
    labels = ", ".join(class_labels(card))
    sign = "predominantly positive" if profile["positive"] >= 0.6 else "predominantly negative" if profile["positive"] <= 0.4 else "state-dependent sign"
    return (
        f"{joint}-{channel}: baseline visited-trajectory mean {fmt(profile['mean'], 1)}, positive fraction {pct(profile['positive'])} ({sign}); "
        f"zeroing in full C11 changes success by {pp(card.get('zero_success_rate_difference_percentage_points'))}, "
        f"and sign reversal by {pp(card.get('sign_flip_success_rate_difference_percentage_points'))}. "
        f"Current frozen label: {labels}. The largest local response is usually driven by {jac['input']} and peaks during {jac['event']}."
    )


def add_conclusions(doc: Document, evidence: Evidence) -> None:
    doc.add_page_break()
    doc.add_heading("4 Established Conclusions in Evidence-Chain Order", level=1)
    doc.add_paragraph(
        "The sequence progresses from whether rolling occurs, through policy response and control action, to failure under intervention. Only frozen endpoint evaluation and frozen interventions adjudicate success, necessity, and sufficiency; atlases, derivatives, and trajectories explain mechanism."
    )
    numbering = create_numbering(doc, ordered=True)
    for text in system_conclusions(evidence):
        add_list_item(doc, text, numbering)
    doc.add_heading("4.1 Current Conclusions for Sixteen K Channels", level=2)
    for joint in JOINTS:
        for channel in CHANNELS:
            add_list_item(doc, channel_conclusion(evidence, joint, channel), numbering)
    old_figure = find_figure(
        OLD_MECHANISM_ROOT,
        "analysis/extended_figures/04_per_joint_sufficiency_necessity_effects.png",
    )
    maybe_add_figure(
        doc,
        old_figure,
        "Figure 3",
        "Per-joint necessity and sufficiency effects in the legacy frozen study",
        "This figure documents prior evidence; the report then completes its causal boundary with zeroing, flipping, dose, and timing interventions on full C11.",
    )


def short_intervention_summary(card: pd.Series) -> str:
    return (
        f"zero {pp(card.get('zero_success_rate_difference_percentage_points'))}; "
        f"flip {pp(card.get('sign_flip_success_rate_difference_percentage_points'))}; "
        f"0.5×/1.5×={pct(card.get('scale_0p5_success_rate'))}/{pct(card.get('scale_1p5_success_rate'))}"
    )


def short_timing_summary(card: pd.Series) -> str:
    return (
        f"static {pp(card.get('static_mean_success_rate_difference_percentage_points'))}; "
        f"permuted {pp(card.get('time_permuted_success_rate_difference_percentage_points'))}"
    )


def add_joint_cards(doc: Document, evidence: Evidence) -> None:
    doc.add_page_break()
    doc.add_heading("5 Sixteen K Mechanism Cards by Joint and Channel", level=1)
    add_callout(
        doc,
        "How to read the cards",
        "Trajectory distribution shows what successful policies used; causal intervention shows whether changing a channel causes failure; local response shows how observations alter K; torque/propagation shows how the change enters the body. All four are required.",
        kind="info",
    )
    for joint_index, joint in enumerate(JOINTS, start=1):
        doc.add_page_break()
        doc.add_heading(f"5.{joint_index} {joint} K1/K2 Mechanism Card", level=2)
        card_rows = []
        process_rows = []
        conclusions = []
        for channel in CHANNELS:
            card = card_row(evidence, joint, channel)
            profile = baseline_profile(evidence, joint, channel)
            jac = dominant_jacobian(evidence, joint, channel)
            physical = physics_summary(evidence, joint, channel)
            propagation = propagation_summary(evidence, joint, channel)
            card_rows.append(
                [
                    channel,
                    f"mean {fmt(profile['mean'], 1)}\n|K| {fmt(profile['abs_mean'], 1)}\npositive {pct(profile['positive'])}",
                    ", ".join(class_labels(card)),
                    short_intervention_summary(card),
                    short_timing_summary(card),
                ]
            )
            sign_direction = "positive" if jac["sign"] > 0.2 else "negative" if jac["sign"] < -0.2 else "bidirectional/state-dependent"
            process_rows.append(
                [
                    channel,
                    f"dominant input {jac['input']}; {jac['event']}; |dK/do| approximately {fmt(jac['abs'], 1)}; {sign_direction}",
                    f"|phi| approximately {fmt(physical['torque'], 2)}; power proxy {fmt(physical['power'], 2)}; total-joint torque saturation {pct(physical['saturation'])}",
                    f"earliest nonlocal propagation: {propagation['family']}/{propagation['role']}, median {fmt(propagation['step'], 0)} steps",
                ]
            )
            conclusions.append(channel_conclusion(evidence, joint, channel))
        add_data_table(
            doc,
            ["Channel", "C11 trajectory distribution", "Frozen classification", "Amplitude/sign intervention", "Timing intervention"],
            card_rows,
            [720, 1600, 2240, 2600, 2200],
            font_size=8.1,
            alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
        )
        add_source_note(doc, "Sources: baseline_joint_profile.csv and per_k_mechanism_cards.csv; all success-rate differences are relative to full C11.")
        add_data_table(
            doc,
            ["Channel", "Local observation to K", "Shapley torque/power", "Closed-loop physical propagation"],
            process_rows,
            [720, 2880, 2880, 2880],
            font_size=8.1,
        )
        add_source_note(doc, "Note: phi*theta_dot is a control-boundary power proxy. Propagation medians use one paired trajectory per condition per seed and are mechanism diagnostics.")
        for index, text in enumerate(conclusions, start=1):
            lead = f"Current {CHANNELS[index-1]} conclusion: "
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(5)
            run = paragraph.add_run(lead)
            set_run_font(run, size=10.3, color=PRESET["dark_blue"], bold=True)
            run = paragraph.add_run(text.split("：", 1)[-1])
            set_run_font(run, size=10.3)


def add_jacobian_section(doc: Document, evidence: Evidence) -> None:
    doc.add_page_break()
    doc.add_heading("6 Local Jacobian: How Observation Changes Rewrite Each K", level=1)
    doc.add_paragraph(
        "The Jacobian does not answer whether a K is necessary; it answers which local-state changes make the policy increase or decrease that K near successfully visited states. "
        "Each joint has only four direct partial derivatives: dK1/d(delta_theta), dK1/d(theta_dot), dK2/d(delta_theta), and dK2/d(theta_dot)."
    )
    full = find_figure(LOCAL_DIR, "figures/J02_full_jacobian_pulse_q3.png")
    derivative = find_figure(LOCAL_DIR, "figures/J03_local_derivative_by_event.png")
    signs = find_figure(LOCAL_DIR, "figures/J04_local_derivative_sign_consistency.png")
    maybe_add_figure(
        doc,
        full,
        "Figure 4",
        "16x16 Jacobian during representative pulse phases",
        "Only eight local 2x2 blocks may be nonzero; blank cross-joint blocks are architectural constraints, not nonsignificant estimates.",
    )
    doc.add_page_break()
    doc.add_heading("6.1 Phase Patterns", level=2)
    maybe_add_figure(
        doc,
        derivative,
        "Figure 5",
        "Local derivatives of 16 K channels under official event bins",
        "The frozen trajectory set has no prelaunch samples, so that column is NA rather than zero; remaining columns compare pulse Q1-Q5 with rolling outside pulses.",
    )
    maybe_add_figure(
        doc,
        signs,
        "Figure 6",
        "Local-derivative sign consistency within the same seed and event bin",
        "This measures whether visited-state derivatives keep the same sign within a bin, not directional consistency across training seeds; cross-seed conclusions require separate comparison of all five seeds.",
    )
    rows = []
    for joint in JOINTS:
        for channel in CHANNELS:
            summary = dominant_jacobian(evidence, joint, channel)
            rows.append(
                [
                    f"{joint}-{channel}",
                    summary["input"],
                    summary["event"],
                    fmt(summary["abs"], 1),
                    pct(summary.get("consistency"), 0),
                    "positive" if summary["sign"] > 0.2 else "negative" if summary["sign"] < -0.2 else "state-dependent",
                ]
            )
    add_data_table(
        doc,
        ["K channel", "Dominant observation", "Strongest phase", "|derivative|", "Sign consistency", "Direction"],
        rows,
        [1440, 1280, 1720, 1400, 1640, 1880],
        font_size=8.1,
        alignments=[WD_ALIGN_PARAGRAPH.CENTER] * 6,
    )
    add_source_note(
        doc,
        "Sign consistency is defined within each seed and event bin as max(positive_fraction, negative_fraction), the fraction of visited-state derivatives sharing a sign; the table reports the cross-seed median after event aggregation. It is not a same-sign vote across five seeds. Prelaunch has no samples and is excluded from dominant-phase comparison.",
    )


def add_physics_section(doc: Document, evidence: Evidence) -> None:
    doc.add_page_break()
    doc.add_heading("7 Shapley Torque, Power, Contact, and Body Propagation", level=1)
    doc.add_paragraph(
        "K1 and K2 add before clipping, so K alone or unclipped terms can double-count or misattribute effects. Exact two-channel Shapley allocation assigns executed clipped torque back to K1 and K2 and reveals when each joint injects, absorbs, or phase-shapes active work across official rolling phases."
    )
    p1 = find_figure(LOCAL_DIR, "figures/P01_shapley_torque_by_event.png")
    p2 = find_figure(LOCAL_DIR, "figures/P02_power_proxy_by_event.png")
    p3 = find_figure(LOCAL_DIR, "figures/P03_saturation_by_event.png")
    maybe_add_figure(
        doc,
        p1,
        "Figure 7",
        "Shapley active-torque magnitude by event phase",
        "Shows the control contribution retained after clipping for every K channel, rather than comparing raw gains alone.",
    )
    doc.add_page_break()
    maybe_add_figure(
        doc,
        p2,
        "Figure 8",
        "Signed power proxy by event phase",
        "Positive values indicate active-work injection into joint motion; negative values indicate absorption/braking. This is a control-boundary proxy, not complete physical energy.",
    )
    maybe_add_figure(
        doc,
        p3,
        "Figure 9",
        "Total joint-control torque saturation rate by event phase",
        "Defined by |u1+u2|>=9, so K1/K2 rows for one joint are identical; it is not an independent per-channel saturation rate. High saturation means increasing K need not increase executed torque.",
    )
    doc.add_heading("7.1 Aggregate All-Joint Saturation", level=2)
    event_rows = []
    for event_bin, group in evidence.event_contact.groupby("event_bin", sort=False):
        event_rows.append(
            [
                EVENT_CN.get(str(event_bin), str(event_bin)),
                fmt(pd.to_numeric(group["step_count"], errors="coerce").median(), 0),
                pct(
                    pd.to_numeric(
                        group["saturation_rate_all_joints"], errors="coerce"
                    ).median(),
                    1,
                ),
                fmt(
                    pd.to_numeric(
                        group["clipped_active_torque_abs_mean"], errors="coerce"
                    ).median(),
                    2,
                ),
                fmt(
                    pd.to_numeric(
                        group["contact_strength_mean"], errors="coerce"
                    ).median(),
                    3,
                ),
            ]
        )
    add_data_table(
        doc,
        ["Event bin", "Median steps per seed", "All-joint saturation rate*", "|clipped torque|", "Contact strength"],
        event_rows,
        [1720, 1800, 1880, 1880, 2080],
        font_size=8.4,
        alignments=[WD_ALIGN_PARAGRAPH.CENTER] * 5,
    )
    add_source_note(
        doc,
        "* saturation_rate_all_joints is the fraction of all joint-steps in an event bin satisfying |u1+u2|>=9, then the median over five training seeds. It is neither the sum of eight joint saturation rates nor a count of saturated joints. Prelaunch is omitted because it has no samples.",
    )
    doc.add_page_break()
    doc.add_heading("7.2 Propagation Order from Local Intervention to Whole-Body Rolling", level=2)
    heat = find_figure(PROPAGATION_DIR, "first_separation_heatmaps.png")
    lag = find_figure(PROPAGATION_DIR, "lag_response_by_joint_distance.png")
    order = find_figure(PROPAGATION_DIR, "contact_rotation_change_order.png")
    maybe_add_figure(
        doc,
        heat,
        "Figure 10",
        "First sustained separation after per-K zeroing/sign reversal",
        "Compares when observations, K, torque, and position at local, adjacent, and distant joints leave baseline; cross-joint K must remain identical at lag 0.",
    )
    maybe_add_figure(
        doc,
        lag,
        "Figure 11",
        "Lag response over different joint distances",
        "Adjacent separation preceding distant separation indicates body-mediated propagation; contact or support changing first suggests amplification by ground constraints.",
    )
    maybe_add_figure(
        doc,
        order,
        "Figure 12",
        "Order of local, contact, and rotation changes",
        "Builds a testable observation -> K -> torque -> pose/contact -> whole-body rotation chain rather than misrepresenting contemporaneous correlation as direct causation.",
    )


def add_causal_interventions(doc: Document, evidence: Evidence) -> None:
    doc.add_page_break()
    doc.add_heading("8 Frozen Causal Interventions: Zeroing, Flipping, Dose, and Timing", level=1)
    doc.add_paragraph(
        "Every channel undergoes six transformations on full C11: zero, 0.5x, 1.5x, sign reversal, static mean, and fixed time permutation. "
        "These separately test presence, correct sign, amplitude window, and dependence on state and timing."
    )
    f1 = find_figure(CAUSAL_DIR, "figures/fig_01_16_channel_effect_heatmap.png")
    f2 = find_figure(CAUSAL_DIR, "figures/fig_02_channel_dose_response.png")
    f3 = find_figure(CAUSAL_DIR, "figures/fig_03_channel_timing_effects.png")
    f5 = find_figure(CAUSAL_DIR, "figures/fig_05_zero_failure_decomposition.png")
    maybe_add_figure(
        doc,
        f1,
        "Figure 13",
        "Success-rate effects for 16 K channels relative to full C11",
        "Compares zeroing, dose, flipping, static replacement, and time permutation; color shows percentage-point change from C11.",
    )
    doc.add_page_break()
    maybe_add_figure(
        doc,
        f2,
        "Figure 14",
        "Per-K dose-response curves",
        "Scales 0, 0.5, 1, and 1.5 show each channel's effective amplitude window; nonmonotonic curves imply joint constraints from clipping and closed-loop phase.",
    )
    maybe_add_figure(
        doc,
        f3,
        "Figure 15",
        "Per-K static-replacement and time-permutation effects",
        "Static mean preserves magnitude but removes state feedback; time permutation preserves marginal distribution but disrupts phase. Together they distinguish value dependence from timing dependence.",
    )
    maybe_add_figure(
        doc,
        f5,
        "Figure 16",
        "Five-part threshold failure decomposition after channel zeroing",
        "Distinguishes whether a channel primarily disrupts launch, rotation, direction, forward displacement, or persistence rather than hiding mechanism differences behind aggregate success.",
    )


def add_joint_pair_section(doc: Document, evidence: Evidence) -> None:
    doc.add_page_break()
    doc.add_heading("9 Joint K1+K2 Necessity and Sufficiency", level=1)
    doc.add_paragraph(
        "Single-channel effects do not automatically imply whole-joint effects. This section treats a joint's K1+K2 as a pair: zero both in C11 to test necessity and transplant both onto C00 to test whether they push the baseline policy into rolling. "
        "The C00 reference is complete in the same Rroll environment, training seeds, and evaluation initial states 20264401-20264420, yielding [18,0,9,0,19]/20 and 46/100 total."
    )
    rows = []
    for joint in JOINTS:
        row = row_for(evidence.pairs, joint=joint)
        if row is None:
            continue
        labels = []
        if as_int(row.get("strong_joint_pair_necessity")):
            labels.append("strongly necessary")
        elif as_int(row.get("joint_pair_necessary_contribution")):
            labels.append("necessary contribution")
        if as_int(row.get("joint_pair_sufficient")):
            labels.append("whole-pair sufficient")
        if as_int(row.get("joint_pair_zero_equivalent_or_redundant")):
            labels.append("approximately equivalent when zeroed")
        rows.append(
            [
                joint,
                pct(row.get("necessity_success_rate")),
                pp(row.get("necessity_success_rate_drop_percentage_points")),
                as_int(row.get("necessity_training_seeds_degraded")),
                pct(row.get("sufficiency_success_rate")),
                pp(row.get("sufficiency_success_rate_gain_percentage_points")),
                ", ".join(labels) or "no frozen label met",
            ]
        )
    add_data_table(
        doc,
        ["Joint", "Whole-pair zero success", "Necessity drop", "Degraded seeds", "Whole-pair transplant success", "Sufficiency gain", "Frozen conclusion"],
        rows,
        [760, 1420, 1260, 1120, 1420, 1260, 2120],
        font_size=7.8,
        alignments=[WD_ALIGN_PARAGRAPH.CENTER] * 7,
    )
    add_callout(
        doc,
        "Sufficiency pairing is complete",
        "Whole-joint K1+K2 sufficiency now uses matched::C00, preserving the same training seeds, Rroll environment, and evaluation seeds 20264401-20264420 as each whole-pair transplant for strict same-initial-state comparison. The 16 legacy single-channel sufficiency labels still come from legacy-C00 pairing within the sealed 59-condition study and must not be presented as matched-C00 results from this study.",
        kind="evidence",
    )
    figure = find_figure(CAUSAL_DIR, "figures/fig_04_joint_pair_effects.png")
    maybe_add_figure(
        doc,
        figure,
        "Figure 17",
        "Whole-pair K1+K2 necessity and sufficiency by joint",
        "Negative bars show loss after whole-pair zeroing in C11; positive bars show gain after whole-pair transplant into C00. They answer different questions.",
    )


def add_mentor_question(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("10 The Research Question the Adviser Actually Asked", level=1)
    add_callout(
        doc,
        "The real question",
        "The goal is neither another proof that one seed can roll nor only two-dimensional K1/K2 response plots. It is to explain how local feedback gains form a reproducible whole-body rolling mechanism under fixed observation/action channels, with intervention-testable explanations for every joint and K.",
        kind="evidence",
    )
    numbering = create_numbering(doc, ordered=True)
    for text in (
        "Policy attribution: is each K controlled primarily by local delta_theta or theta_dot, and how do direction, nonlinearity, and phase vary?",
        "Control attribution: after clipping, how much executed torque comes from K1*delta_theta and K2*theta_dot, and when does each inject, brake, or phase-shape?",
        "Motion chain: how does a local K change sequentially alter its joint, adjacent joints, support/contact, and net body rotation?",
        "Causal necessity: after zeroing or flipping a K, which component of launch, direction, speed, forward motion, or persistence fails first?",
        "Causal sufficiency: can a single channel or whole-joint K1+K2 push a nonrolling baseline into rolling, or is multijoint cooperation required?",
        "Cross-seed robustness: do these patterns persist across five independently learned results from 9201-9205 rather than explaining one attractive trajectory?",
        "Reproducible evidence: can every conclusion be reconstructed from frozen checkpoints, raw trajectories, formulas, tables, figures, and hashes?",
    ):
        add_list_item(doc, text, numbering)
    doc.add_paragraph(
        "The most appropriate thesis framing is therefore how local state-dependent feedback gains emerge through flexible-body and contact feedback as cross-seed robust rolling, "
        "not finding a static K set that makes the robot roll. Static distributions are trajectory descriptions; response rules, phase, torque, and intervention are primary questions."
    )


def unresolved_questions(evidence: Evidence) -> list[str]:
    questions = [
        "Complete physical energy is not exactly decomposed: phi*theta_dot is a control-boundary power proxy and does not integrate all conservative, damping, and contact work over ten physics substeps.",
        "Per-K propagation uses one paired trajectory per condition per seed, sufficient for propagation order but not narrow cross-initial-state intervals; add frozen trajectory repetitions rather than retraining.",
        "Whole-joint matched-C00 same-state sufficiency is complete, but 16 legacy single-channel sufficiency labels remain internally paired in the legacy 59-condition study. Including each single K on initial states 20264401-20264420 requires a separately frozen same-state single-channel transplant matrix.",
        "Current interventions begin at step 0 and cannot fully separate launch from maintenance causal windows; phase-triggered short-window interventions on frozen checkpoints remain necessary.",
        "Higher-order synergy beyond two channels and whole joints is not enumerated; group interactions such as J01+J02 or posterior+midbody may not equal linear sums of single-joint effects.",
        "The local Jacobian is a first-order explanation near visited states; second-order curvature and finite-amplitude counterfactuals may point differently in saturated, turning, or strongly nonlinear regions.",
        "Transfer to hardware or altered friction, mass, and time step is unvalidated; conclusions cover only the frozen physics contract.",
        "Support/contact proxies quantify contact order, but per-entity contact points, normal force, frictional work, and impulse remain unresolved.",
    ]
    for joint in JOINTS:
        for channel in CHANNELS:
            row = card_row(evidence, joint, channel)
            labels = class_labels(row)
            if labels == [CLASS_CN["no_frozen_threshold_label"]]:
                questions.append(
                    f"{joint}-{channel} meets neither strong necessity, necessary contribution, timing-critical, nor equivalence criteria; phase-triggered or grouped interventions must distinguish weak contribution, conditional redundancy, and higher-order synergy."
                )
    return questions


def add_unresolved(doc: Document, evidence: Evidence) -> None:
    doc.add_page_break()
    doc.add_heading("11 Unresolved Questions by Priority", level=1)
    doc.add_paragraph(
        "Unresolved does not mean failed; these questions define the evidence boundary and the smallest additions needed in the next read-only evaluation or frozen intervention."
    )
    numbering = create_numbering(doc, ordered=True)
    for question in unresolved_questions(evidence):
        add_list_item(doc, question, numbering)
    add_callout(
        doc,
        "Stopping rule",
        "Without same-initial-state phase-triggered intervention, exact contact forces, and substep energy, no K should be described as independently sufficient to create rolling, and derivative correlation must not be described as direct cross-joint causation.",
        kind="risk",
    )


def add_evidence_boundaries_and_reproduction(doc: Document, evidence: Evidence) -> None:
    doc.add_page_break()
    doc.add_heading("12 Evidence Boundary and Reproduction Path", level=1)
    doc.add_heading("12.1 Statements Safe to Include in the Thesis", level=2)
    bullets = create_numbering(doc, ordered=False)
    for text in (
        "Under the frozen obs2/action2 contract, Rroll meets the per-seed threshold in 5/5 formal seeds, supporting robust cross-seed reproduction.",
        "The formal actor contains eight local 2-to-2 networks and has a block-diagonal direct Jacobian; cross-joint effects arise from closed-loop physical propagation.",
        "Per-K causal labels follow only frozen rules; channels missing a threshold remain unresolved rather than being relabeled necessary or useless.",
        "Shapley results are post-clipping active-torque contributions and phi*theta_dot is a power proxy, never exact energy.",
        "Whole-joint K1+K2 sufficiency uses matched::C00 paired by seed and initial state; legacy single-channel sufficiency remains explicitly labeled as internally paired in the legacy 59-condition study.",
    ):
        add_list_item(doc, text, bullets)
    doc.add_heading("12.2 Shortest Reproduction Procedure", level=2)
    steps = create_numbering(doc, ordered=True)
    for text in (
        "Verify study_contract.json, study_config.json, SOURCE_MANIFEST.json, matched_c00/MATCHED_C00_COMPLETE.json, and three analysis completion gates.",
        "Use analyze_causal_completion_results.py to read matched::C00 and rebuild per-K causal cards, same-state joint necessity/sufficiency tables, and five-part failure decomposition.",
        "Use analyze_local_actor_jacobian_physics.py to rebuild 2x2 local Jacobians, the 16x16 structural matrix, and Shapley tables, then run the plotting script.",
        "Use analyze_physical_propagation.py on paired trajectories to rebuild first sustained separation, lag response, and contact-rotation order.",
        "Run this script to generate DOCX, then render every page with the documents skill's render_docx.py and inspect/iterate page by page.",
        "Audit final DOCX styles, table DXA geometry, image accessibility, and page count before writing the final report manifest.",
    ):
        add_list_item(doc, text, steps)
    doc.add_heading("12.3 Key Evidence Files", level=2)
    display_paths = [
        ROOT / "study_contract.json",
        MATCHED_C00_ROOT / "MATCHED_C00_COMPLETE.json",
        CAUSAL_DIR / "classification.json",
        CAUSAL_DIR / "per_k_mechanism_cards.csv",
        LOCAL_DIR / "ANALYSIS_MANIFEST.json",
        LOCAL_DIR / "jacobian_local_summary_by_seed_event.csv",
        LOCAL_DIR / "shapley_torque_power_by_seed_event.csv",
        PROPAGATION_DIR / "ANALYSIS_AUDIT.json",
        PROPAGATION_DIR / "per_channel_propagation.csv",
        OLD_MECHANISM_ROOT / "FINAL_MECHANISM_SYNTHESIS.md",
        ATLAS_ROOT / "data" / "analysis_summary.json",
    ]
    rows = [[path.name, str(path), sha256(path)[:16] + "…"] for path in display_paths if path.is_file()]
    add_data_table(
        doc,
        ["File", "Absolute path", "SHA-256 (first 16 characters)"],
        rows,
        [1800, 5760, 1800],
        font_size=7.5,
    )
    add_source_note(doc, "Complete SHA-256 values are written to FINAL_REPORT_MANIFEST.json; the script fails closed if a frozen input changes before or after generation.")
    doc.add_page_break()
    doc.add_heading("Conclusion", level=1)
    doc.add_paragraph(
        "The most defensible conclusion is that rolling is neither a set of large static gains nor an isolated action by one joint; it is a temporally organized feedback field formed by eight local 2-to-2 policies in the body-ground loop. "
        "K1 primarily organizes spatial direction and the curling scaffold; K2 primarily provides dynamic phase shaping, injection, and dissipation. Every joint's specific role requires joint confirmation from per-K interventions, local derivatives, and post-clipping control contributions."
    )
    add_callout(
        doc,
        "Final boundary",
        "This report connects whether rolling occurs, which K structures work, how each K responds to local observations, and how changes propagate into an auditable evidence chain. Remaining uncertainties are listed individually and cannot be overstated as conclusions.",
        kind="evidence",
    )


def add_key_historical_figures(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("Appendix A: Selected Key Historical Figures", level=1)
    doc.add_paragraph(
        "The historical atlas contains 710 figures. This report does not repeat every checkpoint plot and retains only three figures locating formal success, endpoint K distribution, and naturally visited trajectories. The complete atlas remains in its original directory."
    )
    success = find_figure(ATLAS_ROOT, "figures/analysis/C01_terminal_formal_success.png")
    natural = find_figure(ATLAS_ROOT, "figures/analysis/E01_natural_joint_k_distribution.png")
    sign = find_figure(
        OLD_MECHANISM_ROOT,
        "analysis/extended_figures/05_module_c_sign_space_success_heatmap.png",
    )
    maybe_add_figure(
        doc,
        success,
        "Figure A1",
        "Formal frozen endpoint success comparison for R0/Rroll",
        "Answers whether rolling is stable; its evidence level exceeds a training curve or single video.",
    )
    maybe_add_figure(
        doc,
        natural,
        "Figure A2",
        "Natural per-joint K distribution on successful trajectories",
        "Describes the K field actually visited by the policy without misrepresenting frequency as causal necessity.",
    )
    maybe_add_figure(
        doc,
        sign,
        "Figure A3",
        "Frozen-success heat map for K1 spatial sign patterns",
        "Supports the positive-J01/negative-J02-J08 directional scaffold and shows mirrored, all-positive, all-negative, or expanded-positive patterns cannot replace it.",
    )


def build_report(evidence: Evidence, output: Path) -> None:
    if _DOCX_IMPORT_ERROR is not None:
        raise RuntimeError(
            "DOCX generation requires the workspace bundled Python with python-docx; "
            "the current interpreter is suitable only for --validate-only"
        ) from _DOCX_IMPORT_ERROR
    doc = Document()
    configure_document_styles(doc)
    doc.core_properties.title = "From Local Observations to Forward Rolling: Per-Joint K1/K2 Causal-Mechanism Completion"
    doc.core_properties.subject = "Sixteen K channels, local Jacobian, Shapley torque/power, and physical propagation in obs2_roll_repro_v2.1"
    doc.core_properties.author = "Graduate Thesis Project Team"
    doc.core_properties.keywords = "K1,K2,reinforcement learning,rolling,local observations,Jacobian,Shapley,causal intervention"
    doc.core_properties.comments = "Generated read-only from frozen analysis results; narrative_proposal preset + editorial_cover."
    add_cover(doc)
    add_body_section(doc)
    add_toc(doc)
    add_executive_summary(doc, evidence)
    add_contract_and_architecture(doc, evidence)
    add_evidence_inventory(doc, evidence)
    add_project_evidence_timeline(doc, evidence)
    add_conclusions(doc, evidence)
    add_joint_cards(doc, evidence)
    add_jacobian_section(doc, evidence)
    add_physics_section(doc, evidence)
    add_causal_interventions(doc, evidence)
    add_joint_pair_section(doc, evidence)
    add_mentor_question(doc)
    add_unresolved(doc, evidence)
    add_evidence_boundaries_and_reproduction(doc, evidence)
    add_key_historical_figures(doc)
    output.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.NamedTemporaryFile(
        prefix=f".{output.stem}.", suffix=".docx", dir=output.parent, delete=False
    ) as handle:
        temporary = Path(handle.name)
    try:
        doc.save(temporary)
        os.replace(temporary, output)
    finally:
        temporary.unlink(missing_ok=True)


def write_manifest(
    evidence: Evidence,
    output: Path,
    input_hashes: Mapping[str, str],
    manifest_path: Path,
) -> None:
    payload = {
        "schema": "obs2_v2_1_k_causal_completion/final_report_manifest/v1",
        "generated_at_utc": datetime.now(timezone.utc).isoformat(),
        "report": str(output),
        "report_sha256": sha256(output),
        "builder": str(Path(__file__).resolve()),
        "builder_sha256": sha256(Path(__file__).resolve()),
        "design": {
            "preset": PRESET["name"],
            "cover": "editorial_cover",
            "page": "US Letter portrait",
            "margins": "1 inch",
            "usable_width_dxa": 9360,
            "table_indent_dxa": 120,
            "body_font": PRESET["body_font"],
            "east_asia_font": PRESET["east_asia_font"],
        },
        "evidence_contract": {
            "observation": "per-joint [delta_theta, theta_dot]",
            "action": "per-joint [K1, K2]",
            "training_seed_inferential_n": 5,
            "checkpoint": 1500,
            "matched_c00": "same training seed and evaluation seeds 20264401..20264420; 46/100",
            "read_only": True,
        },
        "input_sha256": dict(input_hashes),
        "input_count": len(input_hashes),
        "analysis_gate_summary": {
            "causal_validation_passed": evidence.classification.get("validation", {}).get("passed"),
            "local_jacobian_validation_passed": evidence.local_manifest.get("validation", {}).get("passed"),
            "physical_propagation_status": evidence.propagation_audit.get("status"),
            "physical_inputs_unchanged": evidence.propagation_audit.get("inputs_unchanged"),
            "matched_c00_status": evidence.matched_c00_gate.get("status"),
            "matched_c00_successes": evidence.matched_c00_gate.get("total_successes"),
        },
        "qa_status": "DOCX_BUILT_RENDER_QA_PENDING",
        "qa_instruction": "Render with documents skill render_docx.py; inspect every page at 100%; iterate before delivery.",
    }
    manifest_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    parser.add_argument("--manifest", type=Path, default=DEFAULT_MANIFEST)
    parser.add_argument(
        "--validate-only",
        action="store_true",
        help="Validate completion gates and print a summary without reading/building the report body.",
    )
    args = parser.parse_args()
    evidence = validate_and_load()
    compatibility = compatibility_audit(evidence)
    if args.validate_only:
        print(
            json.dumps(
                {
                    "status": "ready",
                    "causal_rows": len(evidence.cards),
                    "joint_pair_rows": len(evidence.pairs),
                    "jacobian_rows": len(evidence.jacobian),
                    "physics_rows": len(evidence.physics),
                    "propagation_rows": len(evidence.propagation),
                    "compatibility": compatibility,
                },
                ensure_ascii=False,
                indent=2,
            )
        )
        return 0
    output = args.output.resolve()
    manifest = args.manifest.resolve()
    for target in (output, manifest):
        try:
            target.relative_to(ROOT.resolve())
        except ValueError as error:
            raise RuntimeError("Report and manifest must remain inside the study root") from error
    input_hashes_before = {str(path): sha256(path) for path in evidence.source_paths}
    build_report(evidence, output)
    input_hashes_after = {str(path): sha256(path) for path in evidence.source_paths}
    if input_hashes_before != input_hashes_after:
        output.unlink(missing_ok=True)
        raise RuntimeError("Frozen analysis inputs changed while building; report was removed")
    write_manifest(evidence, output, input_hashes_after, manifest)
    print(json.dumps({"status": "built", "report": str(output), "manifest": str(manifest)}, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
